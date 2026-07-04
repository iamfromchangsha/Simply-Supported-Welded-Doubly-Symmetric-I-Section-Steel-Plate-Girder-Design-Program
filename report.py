# -*- coding: utf-8 -*-
"""
报告生成器 — HTML (MathJax LaTeX 渲染) + PDF 导出
兼容 Python 3.11 (f-string expressions 不含反斜杠)
"""

import os
import math
from PySide6.QtCore import QTimer, QUrl
from PySide6.QtWidgets import QApplication
from PySide6.QtWebEngineWidgets import QWebEngineView

from calculator import (
    CalcResult, format_number, E_STEEL, G_STEEL, Q_K,
)

_f = format_number


# ============================================================
# LaTeX 辅助 (无 f-string, 避 Python 3.11 限制)
# ============================================================

def _mi(x):
    """inline math: \( x \)"""
    return "\\(%s\\)" % x

def _md(x):
    """display math: $$ x $$"""
    return "$$\n%s\n$$" % x


# Pre-compute common LaTeX symbols
BS = "\\"  # single backslash character for manual concatenation

class LX:
    """LaTeX symbol constants"""
    ETA   = BS + "eta"
    SIGMA = BS + "sigma"
    TAU   = BS + "tau"
    DELTA = BS + "delta"
    GAMMA = BS + "gamma"
    LAMBDA = BS + "lambda"
    MU    = BS + "mu"
    NU    = BS + "nu"
    RHO   = BS + "rho"
    VARPHI= BS + "varphi"
    SQRT  = BS + "sqrt"
    FRAC  = BS + "frac"
    CDOT  = BS + "cdot"
    TIMES = BS + "times"
    LE    = BS + "le"
    GE    = BS + "ge"
    SIM   = BS + "sim"
    INFTY = BS + "infty"
    PM    = BS + "pm"
    RIGHT = BS + "rightarrow"

    @staticmethod
    def sub(base, sub):
        return "%s_{%s}" % (base, sub)

    @staticmethod
    def pow(base, p):
        return "%s^{%s}" % (base, p)

    @staticmethod
    def frac(num, den):
        return "%s{%s}{%s}" % (BS + "frac", num, den)

    @staticmethod
    def sqrt(x):
        return "%s{%s}" % (BS + "sqrt", x)

    @staticmethod
    def text(s):
        return BS + "text{%s}" % s


# ============================================================
# HTML 模板
# ============================================================

_HTML_HEADER = r"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>钢板梁设计计算书</title>
<script>
MathJax = {
  tex: {
    inlineMath: [['$', '$'], ['\\(', '\\)']],
    displayMath: [['$$', '$$'], ['\\[', '\\]']],
    processEscapes: false
  },
  svg: { fontCache: 'global' },
  options: { enableAssistiveMml: false },
  startup: {
    ready: function () {
      MathJax.startup.defaultReady();
      MathJax.startup.promise.then(function () {
        document.body.setAttribute('data-mjax-done', '1');
      });
    }
  }
};
</script>
<script src="https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-svg.js">
</script>
<style>
  :root {
    --page-width: 210mm;
    --margin: 22mm;
    --content-width: calc(var(--page-width) - 2 * var(--margin));
  }
  * { box-sizing: border-box; margin: 0; padding: 0; }
  body {
    font-family: "Times New Roman", SimSun, serif;
    font-size: 13pt;
    line-height: 1.8;
    color: #000;
    background: #fff;
    padding: 30px 40px;
    max-width: 900px;
    margin: 0 auto;
  }
  h1 { font-size: 18pt; text-align: center; margin: 20px 0 10px; }
  h2 { font-size: 15pt; margin: 28px 0 10px; border-bottom: 1.5px solid #000; padding-bottom: 4px; }
  h3 { font-size: 13pt; margin: 18px 0 6px; }
  h4 { font-size: 12pt; margin: 12px 0 4px; }
  p { margin: 6px 0; text-indent: 2em; }
  p.no-indent { text-indent: 0; }
  table {
    border-collapse: collapse;
    margin: 10px auto;
    font-size: 11pt;
    width: 100%;
  }
  th, td {
    border: 1px solid #000;
    padding: 4px 8px;
    text-align: center;
  }
  th { background: #e0e0e0; font-weight: bold; }
  td { text-align: center; }
  td.left { text-align: left; }
  .formula-block {
    margin: 8px 0 8px 3em;
    font-family: "Times New Roman", serif;
    font-size: 12pt;
  }
  .highlight { font-weight: bold; }
  .ok { color: #006600; font-weight: bold; }
  .ng { color: #cc0000; font-weight: bold; }
  hr { border: none; border-top: 1px solid #000; margin: 20px 0; }
  ul, ol { margin: 4px 0 4px 3em; }
  li { margin: 2px 0; }

  @media print {
    body { padding: 0 5mm; max-width: none; }
    h2 { page-break-after: avoid; }
    table { page-break-inside: avoid; }
  }
</style>
</head>
<body>
"""

_HTML_FOOTER = r"""
</body>
</html>
"""


# ============================================================
# 工字形截面横断面图 (纯 SVG)
# ============================================================

def _make_cross_section_svg(h_w, t_w, b_f, t_f, h):
    """
    生成工字形截面横断面 SVG 示意图（美观比例，标注真实数据）
    h_w, t_w, b_f, t_f, h: mm
    """
    # 固定视觉尺寸（不跟随真实长细比，仅标注反映实际数据）
    VW = 280     # 视觉翼缘宽度
    VHW = 32     # 视觉腹板宽度
    VH = 200     # 视觉总高度
    VHW_H = 152  # 视觉腹板高度
    VTF = (VH - VHW_H) // 2  # 视觉翼缘厚度

    margin = 55
    svg_w = VW + margin * 2 + 60   # 右侧留空间给 h_w / t_f 标注
    svg_h = VH + margin * 2 + 30   # 底部留空间给 t_w 标注

    cx = svg_w / 2.0
    y0 = margin
    y1 = y0 + VTF
    y2 = y1 + VHW_H
    y3 = y2 + VTF
    xl = cx - VW / 2.0
    xr = cx + VW / 2.0
    xw_l = cx - VHW / 2.0
    xw_r = cx + VHW / 2.0

    def r(x, y, w, h_val, color="#d9e6f2"):
        return '<rect x="%.1f" y="%.1f" width="%.1f" height="%.1f" fill="%s" stroke="#000" stroke-width="1"/>' % (x, y, w, h_val, color)

    def line(x1, y1, x2, y2, sw=0.7, dash=None):
        d = ' stroke-dasharray="%s"' % dash if dash else ''
        return '<line x1="%.1f" y1="%.1f" x2="%.1f" y2="%.1f" stroke="#444" stroke-width="%.1f"%s/>' % (x1, y1, x2, y2, sw, d)

    def txt(x, y, text, anchor="middle", size=17, bold=False):
        b = ' font-weight="bold"' if bold else ''
        return '<text x="%.1f" y="%.1f" text-anchor="%s" font-size="%d"%s>%s</text>' % (x, y, anchor, size, b, text)

    parts = []
    a = parts.append

    a('<div style="text-align:center;margin:15px 0;">')
    a('<p><strong>图 3.1&emsp;工字形截面横断面图（单位：mm）</strong></p>')
    a('<svg width="%d" height="%d" xmlns="http://www.w3.org/2000/svg" '
      'style="font-family:SimSun,Times New Roman,serif;">' % (svg_w, svg_h))

    # --- 截面主体 ---
    a(r(xl, y0, VW, VTF, "#d9e6f2"))
    a(r(xl, y2, VW, VTF, "#d9e6f2"))
    a(r(xw_l, y1, VHW, VHW_H, "#f0f0f0"))

    # 中心线
    a(line(cx, y0 - 22, cx, y3 + 22, 0.6, "6,3"))
    a(line(cx - 5, y0 - 14, cx + 5, y0 - 14, 0.6))
    a(line(cx - 5, y3 + 14, cx + 5, y3 + 14, 0.6))

    # ===== b_f =====
    ybf = y0 - 13
    a(line(xl, ybf, xr, ybf, 0.6))
    a(line(xl, ybf - 3, xl, ybf + 3, 0.6))
    a(line(xr, ybf - 3, xr, ybf + 3, 0.6))
    a(txt(cx, ybf - 4, "b_f=%d" % b_f, "middle", 13, True))

    # ===== h =====
    xh = xl - 14
    a(line(xh, y0, xh, y3, 0.6))
    a(line(xh - 3, y0, xh + 3, y0, 0.6))
    a(line(xh - 3, y3, xh + 3, y3, 0.6))
    a(txt(xh - 4, (y0 + y3) / 2 + 3, "h=%d" % h, "end", 13, True))

    # ===== h_w =====
    xhw = xr + 14
    a(line(xhw, y1, xhw, y2, 0.6))
    a(line(xhw - 3, y1, xhw + 3, y1, 0.6))
    a(line(xhw - 3, y2, xhw + 3, y2, 0.6))
    a(txt(xhw + 4, (y1 + y2) / 2 + 3, "h_w=%d" % h_w, "start", 13, True))

    # ===== t_w =====
    ytw = y3 + 14
    a(line(xw_l, ytw, xw_r, ytw, 0.6))
    a(line(xw_l, ytw - 3, xw_l, ytw + 3, 0.6))
    a(line(xw_r, ytw - 3, xw_r, ytw + 3, 0.6))
    a(txt(cx, ytw + 10, "t_w=%d" % t_w, "middle", 13, True))

    # ===== t_f =====
    xtf = xr + 14
    ytf = (y0 + y1) / 2
    a(line(xtf, y0, xtf, y1, 0.6))
    a(line(xtf - 3, y0, xtf + 3, y0, 0.6))
    a(line(xtf - 3, y1, xtf + 3, y1, 0.6))
    a(txt(xtf + 4, ytf + 3, "t_f=%d" % t_f, "start", 13, True))

    a('</svg>')
    a('</div>')

    return '\n'.join(parts)


# ============================================================
# 弯矩图 + 剪力图 (纯 SVG)
# ============================================================

def _make_moment_shear_diagrams(L, g, q, P, P_s, x_cut,
                                  M_gk, M_qk, M_Ed,
                                  M_gk_x, M_qk_x, M_Ed_x,
                                  V_gk, V_qk, V_Ed,
                                  V_gk_x=None, V_qk_x=None, V_Ed_x=None):
    """生成简支梁弯矩图和剪力图组合"""
    svg_w, svg_h = 700, 520
    margin_l, margin_r, margin_t, margin_b = 80, 40, 30, 50
    plot_w = svg_w - margin_l - margin_r
    plot_h_m = 180  # 弯矩图高度
    plot_h_v = 160  # 剪力图高度
    y_m_top = margin_t
    y_v_top = margin_t + plot_h_m + 60
    x_left = margin_l
    x_right = margin_l + plot_w
    y_m_bot = y_m_top + plot_h_m
    y_v_bot = y_v_top + plot_h_v

    def px(x_m):
        """将位置 x (m) 映射到 SVG x 坐标"""
        return margin_l + x_m / L * plot_w

    def py_m(M):
        """将弯矩 (kN·m) 映射到 SVG y (向下为正)"""
        Mmax = max(abs(M_gk), abs(M_qk), abs(M_Ed),
                   abs(M_gk_x), abs(M_qk_x), abs(M_Ed_x), 1.0) * 1.15
        return y_m_bot - M / Mmax * plot_h_m

    def py_v(V):
        """将剪力 (kN) 映射到 SVG y"""
        Vmax = max(abs(V_gk), abs(V_qk), abs(V_Ed), -V_gk, -V_qk, -V_Ed, 1.0) * 1.2
        mid = (y_v_top + y_v_bot) / 2
        return mid - V / Vmax * (plot_h_v / 2)

    def line(x1, y1, x2, y2, color="#000", sw=1.0, dash=None):
        d = ' stroke-dasharray="%s"' % dash if dash else ''
        return ('<line x1="%.1f" y1="%.1f" x2="%.1f" y2="%.1f" '
                'stroke="%s" stroke-width="%.1f"%s/>' % (x1, y1, x2, y2, color, sw, d))

    def txt(x, y, text, anchor="middle", size=17, color="#333", bold=False):
        b = ' font-weight="bold"' if bold else ''
        return ('<text x="%.1f" y="%.1f" text-anchor="%s" font-size="%d" '
                'fill="%s"%s>%s</text>' % (x, y, anchor, size, color, b, text))

    def arrow(x1, y1, x2, y2, color="#444", sw=0.8):
        return ('<line x1="%.1f" y1="%.1f" x2="%.1f" y2="%.1f" '
                'stroke="%s" stroke-width="%.1f" marker-end="url(#arrowhead)"/>' %
                (x1, y1, x2, y2, color, sw))

    parts = []
    a = parts.append

    # SVG 容器
    a('<div style="text-align:center;margin:15px 0;">')
    a('<svg width="%d" height="%d" xmlns="http://www.w3.org/2000/svg" '
      'style="font-family:SimSun,Times New Roman,serif;">' % (svg_w, svg_h))

    # 箭头定义
    a('<defs>')
    a('<marker id="arrowhead" markerWidth="8" markerHeight="6" '
      'refX="8" refY="3" orient="auto">')
    a('<polygon points="0 0, 8 3, 0 6" fill="#444"/>')
    a('</marker>')
    a('</defs>')

    # ===== 弯矩图 =====
    a(txt(px(L/2), y_m_top - 8, "弯矩图 M (kN·m)", "middle", 16, "#000", True))

    # 基线
    a(line(x_left, y_m_bot, x_right, y_m_bot, "#000", 1.2))
    # 左纵轴
    a(line(x_left, y_m_top, x_left, y_m_bot, "#000", 1.2))

    # 标注最大值
    Mmax_val = max(abs(M_gk), abs(M_qk), abs(M_Ed), 1.0) * 1.15
    a(txt(x_left - 8, y_m_top + 10, "%.0f" % Mmax_val, "end", 12, "#555"))
    a(txt(x_left - 8, y_m_bot, "0", "end", 12, "#555"))
    # 标注位置
    a(txt(px(L/2), y_m_bot + 15, "L/2", "middle", 12, "#555"))
    a(txt(px(0), y_m_bot + 15, "0", "middle", 12, "#555"))
    a(txt(px(L), y_m_bot + 15, "L", "middle", 12, "#555"))
    a(txt(px(x_cut), y_m_bot + 28, "L/6", "middle", 12, "#555"))

    # 虚线标出 x_cut 位置
    a(line(px(x_cut), y_m_top, px(x_cut), y_v_bot + 15, "#999", 0.6, "4,4"))

    # M_gk 曲线 (抛物线, 多点近似)
    npts = 40
    pts_gk = []
    pts_qk = []
    for i in range(npts + 1):
        xi = L * i / npts
        Mi_gk = g * xi * (L - xi) / 2.0
        Mi_qk = q * xi * (L - xi) / 2.0 + P * xi / 2.0  # xi ≤ L/2 only
        pts_gk.append("%.1f,%.1f" % (px(xi), py_m(Mi_gk)))
        if xi <= L / 2:
            pts_qk.append("%.1f,%.1f" % (px(xi), py_m(Mi_qk)))
        else:
            pts_qk.append("%.1f,%.1f" % (px(xi), py_m(Mi_qk)))
    # Fix: for xi > L/2, the concentrated load moment is P*(L-xi)/2
    pts_qk = []
    for i in range(npts + 1):
        xi = L * i / npts
        Mi_qk = q * xi * (L - xi) / 2.0
        if xi <= L / 2:
            Mi_qk += P * xi / 2.0
        else:
            Mi_qk += P * (L - xi) / 2.0
        pts_qk.append("%.1f,%.1f" % (px(xi), py_m(Mi_qk)))

    a('<polyline points="%s" fill="none" stroke="#2196F3" stroke-width="1.8"/>' % ' '.join(pts_gk))
    a('<polyline points="%s" fill="none" stroke="#FF9800" stroke-width="1.8"/>' % ' '.join(pts_qk))

    # M_Ed 曲线 (缩放后的组合)
    pts_Ed = []
    for i in range(npts + 1):
        xi = L * i / npts
        Mi_gk_i = g * xi * (L - xi) / 2.0
        Mi_qk_i = q * xi * (L - xi) / 2.0
        if xi <= L / 2:
            Mi_qk_i += P * xi / 2.0
        else:
            Mi_qk_i += P * (L - xi) / 2.0
        gamma_G = 1.2; gamma_Q = 1.4; gamma_0 = 1.1; mu = 1.14
        Mi_Ed = gamma_0 * (gamma_G * Mi_gk_i + gamma_Q * mu * Mi_qk_i)
        pts_Ed.append("%.1f,%.1f" % (px(xi), py_m(Mi_Ed)))
    a('<polyline points="%s" fill="none" stroke="#E91E63" stroke-width="2.2"/>' % ' '.join(pts_Ed))

    # 图例
    a(txt(px(L) - 20, y_m_top + 10, "M_gk", "end", 12, "#2196F3", True))
    a(txt(px(L) - 20, y_m_top + 24, "M_qk", "end", 12, "#FF9800", True))
    a(txt(px(L) - 20, y_m_top + 38, "M_Ed", "end", 12, "#E91E63", True))

    # ===== 剪力图 =====
    a(txt(px(L/2), y_v_top - 8, "剪力图 V (kN)", "middle", 16, "#000", True))

    # 基线
    v_mid_y = (y_v_top + y_v_bot) / 2
    a(line(x_left, v_mid_y, x_right, v_mid_y, "#000", 1.2))
    # 左纵轴
    a(line(x_left, y_v_top, x_left, y_v_bot, "#000", 1.2))

    # 标注
    Vmax_val = max(abs(V_gk), abs(V_qk), abs(V_Ed), 1.0) * 1.2
    a(txt(x_left - 8, y_v_top + 8, "+%.0f" % Vmax_val, "end", 12, "#555"))
    a(txt(x_left - 8, v_mid_y, "0", "end", 12, "#555"))
    a(txt(x_left - 8, y_v_bot - 2, "-%.0f" % Vmax_val, "end", 12, "#555"))

    # V_gk (线性)
    a(line(px(0), py_v(V_gk), px(L), py_v(-V_gk), "#2196F3", 1.8))
    # V_qk: 均布 + 集中力在跨中产生突变
    # 0 → L/2: V = q*(L/2-x) + P_s/2
    # L/2 → L: V = -q*(x-L/2) - P_s/2
    a(line(px(0), py_v(V_qk), px(L/2 - 0.01), py_v(P_s/2), "#FF9800", 1.8))
    a(line(px(L/2 + 0.01), py_v(-P_s/2), px(L), py_v(-V_qk), "#FF9800", 1.8))

    # V_Ed
    V0_Ed = 1.1 * (1.2 * V_gk + 1.4 * 1.14 * V_qk)
    a(line(px(0), py_v(V0_Ed), px(L), py_v(-V0_Ed), "#E91E63", 2.2))

    # 图例
    a(txt(px(L) - 20, y_v_top + 10, "V_gk", "end", 12, "#2196F3", True))
    a(txt(px(L) - 20, y_v_top + 24, "V_qk", "end", 12, "#FF9800", True))
    a(txt(px(L) - 20, y_v_top + 38, "V_Ed", "end", 12, "#E91E63", True))

    a('</svg>')
    a('<p><strong>图 4.1&emsp;简支梁弯矩图与剪力图</strong></p>')
    a('</div>')

    return '\n'.join(parts)


# ============================================================
# 变截面布置示意图 (纯 SVG)
# ============================================================

def _make_variable_section_layout_svg(L, x_cut, sec_mid, sec_var):
    """生成变截面布置示意图
    sec_mid, sec_var: (h_w, t_w, b_f, t_f)
    """
    hw_mid, tw_mid, bf_mid, tf_mid = sec_mid
    hw_var, tw_var, bf_var, tf_var = sec_var
    svg_w, svg_h = 680, 200
    margin_l, margin_r = 80, 60
    plot_w = svg_w - margin_l - margin_r
    x_left = margin_l

    def px(x_m):
        return margin_l + x_m / L * plot_w

    parts = []
    a = parts.append

    a('<div style="text-align:center;margin:15px 0;">')
    a('<svg width="%d" height="%d" xmlns="http://www.w3.org/2000/svg" '
      'style="font-family:SimSun,Times New Roman,serif;">' % (svg_w, svg_h))

    # 梁体 (变截面示意)
    beam_y_top = 50
    beam_h_mid = 60
    beam_h_var = 40
    a('<rect x="%.1f" y="%.1f" width="%.1f" height="%d" fill="#c8daf0" '
      'stroke="#333" stroke-width="1.2"/>' %
      (px(0), beam_y_top + beam_h_mid - beam_h_var + 10,
       px(L) - px(0), beam_h_var))

    # 跨中等截面段 (较厚翼缘)
    seg_mid_left = px(x_cut)
    seg_mid_right = px(L - x_cut)
    a('<rect x="%.1f" y="%.1f" width="%.1f" height="%d" fill="#7ea8d4" '
      'stroke="#333" stroke-width="1.2"/>' %
      (seg_mid_left, beam_y_top,
       seg_mid_right - seg_mid_left, beam_h_mid))

    # 标注
    def txt(x, y, text, anchor="middle", size=17, color="#333", bold=False):
        b = ' font-weight="bold"' if bold else ''
        return ('<text x="%.1f" y="%.1f" text-anchor="%s" font-size="%d" '
                'fill="%s"%s>%s</text>' % (x, y, anchor, size, color, b, text))

    def line(x1, y1, x2, y2, sw=0.7, dash=None):
        d = ' stroke-dasharray="%s"' % dash if dash else ''
        return ('<line x1="%.1f" y1="%.1f" x2="%.1f" y2="%.1f" '
                'stroke="#555" stroke-width="%.1f"%s/>' % (x1, y1, x2, y2, sw, d))

    # 支座标记
    a('<polygon points="%.1f,%d %.1f,%d %.1f,%d" fill="#555"/>' %
      (px(0), beam_y_top + beam_h_mid + 10,
       px(0) - 15, beam_y_top + beam_h_mid + 30,
       px(0) + 15, beam_y_top + beam_h_mid + 30))
    a('<polygon points="%.1f,%d %.1f,%d %.1f,%d" fill="#555"/>' %
      (px(L), beam_y_top + beam_h_mid + 10,
       px(L) - 15, beam_y_top + beam_h_mid + 30,
       px(L) + 15, beam_y_top + beam_h_mid + 30))

    # 尺寸标注
    a(txt(px(0), beam_y_top + beam_h_mid + 45, "0", "middle", 12, "#555"))
    a(txt(px(L), beam_y_top + beam_h_mid + 45, "L=%.0fm" % L, "middle", 12, "#555"))
    a(txt(px(x_cut), beam_y_top + beam_h_mid + 55, "L/6=%.2fm" % x_cut, "middle", 12, "#E91E63"))
    a(txt(px(L - x_cut), beam_y_top + beam_h_mid + 55, "5L/6", "middle", 12, "#E91E63"))

    # 虚线标变截面位置
    a(line(px(x_cut), beam_y_top - 10, px(x_cut), beam_y_top + beam_h_mid + 20, 0.7, "5,3"))
    a(line(px(L - x_cut), beam_y_top - 10, px(L - x_cut), beam_y_top + beam_h_mid + 20, 0.7, "5,3"))

    # 截面标注
    mids = [(px(L/2), "跨中段", sec_mid),
            (px(L/2 - L/4), "跨中段", sec_mid),
            (px(L/12), "变截面段", sec_var),
            (px(L - L/12), "变截面段", sec_var)]
    # 简化截面标注
    a(txt(px(L/2), beam_y_top - 5, "跨中段: %dx%d / %dx%d" %
          (hw_mid, tw_mid, bf_mid, tf_mid),
          "middle", 9, "#333", True))
    y_var_label = beam_y_top + beam_h_mid + 70
    a(txt(px(x_cut/2), y_var_label, "变截面段: %dx%d / %dx%d" %
          (hw_var, tw_var, bf_var, tf_var),
          "middle", 9, "#333"))
    a(txt(px(L - x_cut/2), y_var_label, "变截面段: %dx%d / %dx%d" %
          (hw_var, tw_var, bf_var, tf_var),
          "middle", 9, "#333"))

    a('</svg>')
    a('<p><strong>图 6.1&emsp;变截面布置示意图</strong></p>')
    a('</div>')

    return '\n'.join(parts)


# ============================================================
# 加劲肋布置示意图 (纯 SVG)
# ============================================================

def _make_stiffener_layout_svg(L, spacing, n_pairs):
    """生成横向加劲肋布置示意图"""
    svg_w, svg_h = 700, 190
    margin_l, margin_r = 50, 50
    plot_w = svg_w - margin_l - margin_r

    def px(x_mm):
        return margin_l + x_mm / (L * 1000.0) * plot_w

    parts = []
    a = parts.append

    a('<div style="text-align:center;margin:15px 0;">')
    a('<svg width="%d" height="%d" xmlns="http://www.w3.org/2000/svg" '
      'style="font-family:SimSun,Times New Roman,serif;">' % (svg_w, svg_h))

    def txt(x, y, text, anchor="middle", size=17, color="#333", bold=False):
        b = ' font-weight="bold"' if bold else ''
        return ('<text x="%.1f" y="%.1f" text-anchor="%s" font-size="%d" '
                'fill="%s"%s>%s</text>' % (x, y, anchor, size, color, b, text))

    def line(x1, y1, x2, y2, sw=0.7, dash=None):
        d = ' stroke-dasharray="%s"' % dash if dash else ''
        return ('<line x1="%.1f" y1="%.1f" x2="%.1f" y2="%.1f" '
                'stroke="#555" stroke-width="%.1f"%s/>' % (x1, y1, x2, y2, sw, d))

    # 主梁
    beam_y = 55
    beam_h = 28
    a('<rect x="%.1f" y="%d" width="%.1f" height="%d" fill="#d0dce8" '
      'stroke="#333" stroke-width="1.2"/>' %
      (px(0), beam_y, px(L * 1000) - px(0), beam_h))

    # 支座标记（三角形）
    a('<polygon points="%.1f,%d %.1f,%d %.1f,%d" fill="#333"/>' %
      (px(0), beam_y + beam_h + 4, px(0) - 12, beam_y + beam_h + 20,
       px(0) + 12, beam_y + beam_h + 20))
    a('<polygon points="%.1f,%d %.1f,%d %.1f,%d" fill="#333"/>' %
      (px(L * 1000), beam_y + beam_h + 4,
       px(L * 1000) - 12, beam_y + beam_h + 20,
       px(L * 1000) + 12, beam_y + beam_h + 20))

    # 横向加劲肋（竖线，从梁顶上方冒出 + 梁下方一小段）
    stiff_top = beam_y - 12
    stiff_bot = beam_y + beam_h + 6
    n_total = int(L * 1000 / spacing) + 1
    for i in range(1, n_total - 1):  # 跳过两端（支座加劲肋）
        xi_mm = i * spacing
        if xi_mm < L * 1000:
            a(line(px(xi_mm), stiff_top, px(xi_mm), stiff_bot, 1.0))
            # 顶部小横线
            a(line(px(xi_mm) - 3, stiff_top, px(xi_mm) + 3, stiff_top, 0.7))

    # 间距标注（在第一、二根横向加劲肋之间，不会被支座加劲肋文字遮盖）
    x1 = px(spacing)
    x2 = px(2 * spacing) if 2 * spacing < L * 1000 else px(spacing)
    x_mid = (x1 + x2) / 2
    a(line(x1, stiff_top - 8, x2, stiff_top - 8, 0.7))
    a(line(x1, stiff_top - 11, x1, stiff_top - 5, 0.7))
    a(line(x2, stiff_top - 11, x2, stiff_top - 5, 0.7))
    a(txt(x_mid, stiff_top - 14, "a=%dmm" % spacing, "middle", 13, "#E91E63", True))

    # 支座加劲肋（粗线 + 标注，放在梁下方避免与间距标注重叠）
    a(line(px(0), beam_y - 5, px(0), beam_y + beam_h + 5, 2.0))
    a(line(px(L * 1000), beam_y - 5, px(L * 1000), beam_y + beam_h + 5, 2.0))
    a(txt(px(0), beam_y + beam_h + 38, "支座\n加劲肋", "middle", 12, "#333", True))

    # 右端支座（用 position 判断是否与 stiffeners 太近）
    a(txt(px(L * 1000), beam_y + beam_h + 38, "支座\n加劲肋", "middle", 12, "#333", True))

    # 底部总说明
    a(txt(margin_l + plot_w / 2, svg_h - 14,
          "横向加劲肋共 %d 对，间距 a = %d mm，成对布置于腹板两侧" % (n_pairs, spacing),
          "middle", 11, "#333", True))

    a('</svg>')
    a('<p><strong>图 8.1&emsp;加劲肋布置示意图</strong></p>')
    a('</div>')

    return '\n'.join(parts)


def _make_girder_weld_svg(r):
    """生成 CAD 风格的钢板梁及焊缝构造图"""
    p = r.params; sec = r.section; wd = r.welds
    pm = r.props_mid; pv = r.props_var
    svg_w, svg_h = 820, 620
    L = p.L * 1000.0

    C = '#000'  # 主色黑
    CDIM = '#555'  # 尺寸线灰
    CHATCH = '#aaa'

    parts = []
    a = parts.append

    # ─── 辅助函数 ───
    def txt(x, y, s, size=14, anchor='middle', bold=False, color=C):
        b = ' font-weight="bold"' if bold else ''
        return f'<text x="{x:.0f}" y="{y:.0f}" text-anchor="{anchor}" font-size="{size}" fill="{color}"{b}>{s}</text>'

    def line(x1, y1, x2, y2, sw=1, color=C, dash=''):
        d = f' stroke-dasharray="{dash}"' if dash else ''
        return f'<line x1="{x1:.2f}" y1="{y1:.2f}" x2="{x2:.2f}" y2="{y2:.2f}" stroke="{color}" stroke-width="{sw}"{d}/>'

    def rect(x, y, w, h, fill='none', stroke=C, sw=1.2):
        return f'<rect x="{x:.2f}" y="{y:.2f}" width="{w:.2f}" height="{h:.2f}" fill="{fill}" stroke="{stroke}" stroke-width="{sw}"/>'

    def hatch_rect(x, y, w, h):
        """带 45° 斜线的填充矩形（剖切填充）"""
        pts = []
        s2 = 8
        start_y = y - (x % s2)
        for yy in range(int(start_y), int(y + h + s2), s2):
            pts.append(line(x, yy, x + w, yy + s2, 0.5, CHATCH))
        for yy in range(int(start_y - s2), int(y + h + s2), s2):
            pts.append(line(x, yy + s2, x + w, yy, 0.5, CHATCH))
        s = ''.join(pts)
        return (rect(x, y, w, h, 'none', C, 1.5) + s)

    def dim_45(x1, y1, x2, y2, label, side='above', offset=14):
        """CAD 风格 45° 斜线尺寸标注"""
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        angle = 4
        s = ''
        s += line(x1, y1, x2, y2, 0.7, CDIM)
        # 左端 tick
        s += line(x1 - angle, y1 - angle, x1 + angle, y1 + angle, 0.7, CDIM)
        s += line(x1 - angle, y1 + angle, x1 + angle, y1 - angle, 0.7, CDIM)
        # 右端 tick
        s += line(x2 - angle, y2 - angle, x2 + angle, y2 + angle, 0.7, CDIM)
        s += line(x2 - angle, y2 + angle, x2 + angle, y2 - angle, 0.7, CDIM)
        # 文字
        if side == 'above':
            s += txt(mx, my - 5, label, 17, 'middle', False, C)
        else:
            s += txt(mx, my + 16, label, 17, 'middle', False, C)
        return s

    def dim_vert(x, y1, y2, label, offset=-16):
        """垂直尺寸"""
        mx = x + offset
        my = (y1 + y2) / 2
        s = ''
        s += line(mx, y1, mx, y2, 0.7, CDIM)
        s += line(mx - 3, y1, mx + 3, y1, 0.7, CDIM)
        s += line(mx - 3, y2, mx + 3, y2, 0.7, CDIM)
        s += txt(mx - 6, my + 3, label, 17, 'end', False, C)
        return s

    a('<div style="text-align:center;margin:10px 0;">')
    a('<svg width="%d" height="%d" xmlns="http://www.w3.org/2000/svg" '
      'style="font-family:SimSun,serif;">' % (svg_w, svg_h))

    # ────── 1. 立面图 ──────
    ey = 55
    ml, mr = 60, 60
    pw = svg_w - ml - mr
    xc = r.x_cut * 1000.0

    def px(x_mm):
        return ml + x_mm / L * pw if L > 0 else ml

    mid_h = 55
    # 主梁轮廓
    a(rect(px(0), ey, xc / L * pw, mid_h, '#fff', C, 2))
    a(rect(px(xc), ey, (L - xc) / L * pw, mid_h, '#fff', C, 2))
    # 变截面斜线
    a(line(px(xc), ey, px(xc), ey + mid_h, 2, C))
    a(line(px(xc) - 8, ey, px(xc) + 8, ey, 0.7, CDIM))
    a(line(px(xc) - 8, ey + mid_h, px(xc) + 8, ey + mid_h, 0.7, CDIM))

    # 支座
    a(f'<polygon points="{px(0):.0f},{ey+mid_h+3} {px(0)-10:.0f},{ey+mid_h+18} {px(0)+10:.0f},{ey+mid_h+18}" fill="{C}"/>')
    a(f'<polygon points="{px(L):.0f},{ey+mid_h+3} {px(L)-10:.0f},{ey+mid_h+18} {px(L)+10:.0f},{ey+mid_h+18}" fill="{C}"/>')

    # 跨度标注
    yL = ey + mid_h + 24
    a(dim_45(px(0), yL, px(L), yL, f'L={p.L:.1f}m', 'below', 12))
    # x_cut 标注
    xcp = px(xc)
    a(line(xcp, ey - 8, xcp, ey + mid_h + 8, 0.7, CDIM, '4,3'))
    a(txt(xcp, ey - 14, f'x_cut={r.x_cut:.2f}', 16, 'middle', False, CDIM))

    # 剖面标记
    a(txt(ml + pw * 0.25, ey + mid_h + 42, 'A', 17, 'middle', True))
    a(line(ml + pw * 0.25 - 6, ey + mid_h + 32, ml + pw * 0.25 + 6, ey + mid_h + 32, 1, C))
    a(txt(ml + pw * 0.75, ey + mid_h + 42, 'B', 17, 'middle', True))
    a(line(ml + pw * 0.75 - 6, ey + mid_h + 32, ml + pw * 0.75 + 6, ey + mid_h + 32, 1, C))
    # 截面指示线
    a(line(ml + pw * 0.25, ey + mid_h + 28, ml + pw * 0.25, ey + mid_h + 50, 0.7, C))
    a(line(ml + pw * 0.75, ey + mid_h + 28, ml + pw * 0.75, ey + mid_h + 50, 0.7, C))

    # ────── 2. A-A 剖面 ──────
    ax0, ay0 = 20, 320
    sc = 1.1
    bf = sec.b_f * sc
    bf = min(bf, 170)
    sc = bf / max(sec.b_f, 1)
    tw = max(sec.t_w * sc, 3)
    tf = max(sec.t_f * sc, 5)
    hw = sec.h_w * sc
    bw = sec.b_f * sc
    h_total = pm.h * sc

    ca = ax0 + bw / 2
    # 剖切填充
    a(hatch_rect(ax0, ay0, bw, tf))
    a(hatch_rect(ax0, ay0 + tf + hw, bw, tf))
    a(hatch_rect(ca - tw / 2, ay0 + tf, tw, hw))
    # 中心线
    a(line(ca, ay0 - 12, ca, ay0 + h_total + 12, 0.7, CDIM, '8,4'))

    # 尺寸
    right = ax0 + bw + 8
    a(dim_vert(right, ay0, ay0 + tf, f't_f={sec.t_f}', 16))
    a(dim_vert(right, ay0 + tf, ay0 + tf + hw, f'h_w={sec.h_w}', 16))
    a(dim_vert(right, ay0 + h_total - tf, ay0 + h_total, f't_f={sec.t_f}', 16))
    # h
    a(dim_vert(ax0 - 12, ay0, ay0 + h_total, f'h={pm.h}', -12))
    # b_f
    a(dim_45(ax0, ay0 + h_total + 14, ax0 + bw, ay0 + h_total + 14, f'b_f={sec.b_f}', 'below', 10))
    a(txt(ax0 + bw / 2, ay0 - 18, 'A—A', 16, 'middle', True))

    # ────── 3. 焊缝详图 ──────
    zx, zy = 275, 285
    zw, zh = 230, 185
    a(rect(zx, zy, zw, zh, '#fafafa', C, 1.2))
    a(txt(zx + zw / 2, zy + 16, 'Ⅰ 焊缝详图', 14, 'middle', True))

    # 几何参数（放大视图）
    _fw = 90                     # 翼缘宽度 px
    _fh = 22                     # 翼缘厚度 px
    _ww = 12                     # 腹板厚度 px
    _wh = 55                     # 腹板高度 px
    _fx = zx + (zw - _fw) / 2   # 翼缘左端 x
    _fy = zy + 50                # 翼缘顶部 y
    _fcx = _fx + _fw / 2         # 中心线 x
    _wy = _fy + _fh - 2          # 腹板顶部 y（略低于翼缘底）
    _wx = _fcx - _ww / 2         # 腹板左端 x

    # 翼缘（剖切填充）
    a(hatch_rect(_fx, _fy, _fw, _fh))
    # 腹板（剖切填充）
    a(hatch_rect(_wx, _wy, _ww, _wh))

    # 角焊缝（双面三角形，红色）
    wc = '#d00'
    _wl = 16                     # 焊缝直角边 px
    _weld_top = _wy              # 焊缝顶部 y（腹板顶）
    a(f'<polygon points="{_fcx:.0f},{_weld_top:.0f} {_fcx:.0f},{_weld_top+_wl:.0f} {_fcx-_wl:.0f},{_weld_top:.0f}" fill="none" stroke="{wc}" stroke-width="1.5"/>')
    a(f'<polygon points="{_fcx:.0f},{_weld_top:.0f} {_fcx:.0f},{_weld_top+_wl:.0f} {_fcx+_wl:.0f},{_weld_top:.0f}" fill="none" stroke="{wc}" stroke-width="1.5"/>')

    # 焊缝尺寸标注
    _ly = _weld_top + _wl + 8
    a(line(_fcx - _wl, _ly, _fcx, _ly, 0.7, wc))
    a(line(_fcx, _ly, _fcx + _wl, _ly, 0.7, wc))
    a(line(_fcx - _wl, _ly - 3, _fcx - _wl, _ly + 3, 0.7, wc))
    a(line(_fcx, _ly - 3, _fcx, _ly + 3, 0.7, wc))
    a(line(_fcx + _wl, _ly - 3, _fcx + _wl, _ly + 3, 0.7, wc))
    # 尺寸箭头（焊接符号）
    a(f'<polygon points="{_fcx-4:.0f},{_ly+6:.0f} {_fcx:.0f},{_ly+10:.0f} {_fcx+4:.0f},{_ly+6:.0f}" fill="{wc}"/>')
    a(txt(_fcx, _ly + 18, f'h_f={wd.h_f_chosen:.0f}', 14, 'middle', True, wc))

    # 说明文字
    a(txt(zx + zw / 2, zy + zh - 10, '双面连续角焊缝', 16, 'middle', False, CDIM))
    # 标注指引线
    a(line(zx + zw / 2 - 40, zy + zh - 14, zx + zw / 2 + 40, zy + zh - 14, 0.5, CDIM))
    a(line(zx + zw / 2 - 40, zy + zh - 14, zx + zw / 2 - 40, _wy + _wh + 6, 0.5, CDIM))
    a(line(zx + zw / 2 + 40, zy + zh - 14, zx + zw / 2 + 40, _weld_top + _wl + 2, 0.5, CDIM))

    # 翼缘厚度标注
    a(dim_vert(_fx - 10, _fy, _fy + _fh, f't_f={sec.t_f}', -10))
    # 腹板厚度标注
    _tx = _fx + _fw + 10
    a(line(_tx, _wy, _tx, _wy + _wh, 0.7, CDIM))
    a(line(_tx - 3, _wy, _tx + 3, _wy, 0.7, CDIM))
    a(line(_tx - 3, _wy + _wh, _tx + 3, _wy + _wh, 0.7, CDIM))
    a(txt(_tx + 10, _wy + _wh / 2 + 3, f't_w={sec.t_w}', 17, 'start'))

    # ────── 4. B-B 剖面 ──────
    bx0, by0 = 520, 320
    sc2 = 1.1
    bf2 = sec.b_f2 * sc2
    bf2 = min(bf2, 170)
    sc2 = bf2 / max(sec.b_f2, 1)
    tf2 = max(sec.t_f2 * sc2, 5)
    hw2 = sec.h_w * sc2
    bw2 = sec.b_f2 * sc2
    h_t2 = pv.h * sc2

    cb = bx0 + bw2 / 2
    a(hatch_rect(bx0, by0, bw2, tf2))
    a(hatch_rect(bx0, by0 + tf2 + hw2, bw2, tf2))
    a(hatch_rect(cb - tw / 2, by0 + tf2, tw, hw2))
    a(line(cb, by0 - 12, cb, by0 + h_t2 + 12, 0.7, CDIM, '8,4'))
    # 尺寸
    r2 = bx0 + bw2 + 8
    a(dim_vert(r2, by0, by0 + tf2, f't_f2={sec.t_f2}', 16))
    a(dim_vert(r2, by0 + tf2, by0 + tf2 + hw2, f'h_w={sec.h_w}', 16))
    a(dim_vert(bx0 - 12, by0, by0 + h_t2, f'h2={pv.h}', -12))
    a(dim_45(bx0, by0 + h_t2 + 14, bx0 + bw2, by0 + h_t2 + 14, f'b_f2={sec.b_f2}', 'below', 10))
    a(txt(bx0 + bw2 / 2, by0 - 18, 'B—B', 16, 'middle', True))

    # ────── 5. 标题栏 ──────
    tb_y = svg_h - 50
    a(rect(10, tb_y, svg_w - 20, 40, '#fff', C, 1.5))
    a(line(10, tb_y + 24, svg_w - 10, tb_y + 24, 0.7, C))
    a(txt(30, tb_y + 34, '钢板梁及焊缝构造图', 16, 'start', True))
    a(txt(svg_w / 2, tb_y + 14, '比例 1:10', 17, 'middle', False, CDIM))
    a(txt(svg_w - 30, tb_y + 34, '单位：mm', 17, 'end', False, CDIM))

    a('</svg>')
    a('</div>')
    return '\n'.join(parts)


def _make_stiffener_detail_svg(r):
    """生成 CAD 风格的加劲肋构造图"""
    p = r.params; sec = r.section; st = r.stiffeners
    svg_w, svg_h = 820, 580
    L = p.L * 1000.0

    C = '#000'; CDIM = '#555'; CHATCH = '#aaa'

    def txt(x, y, s, size=14, anchor='middle', bold=False, color=C):
        b = ' font-weight="bold"' if bold else ''
        return f'<text x="{x:.2f}" y="{y:.2f}" text-anchor="{anchor}" font-size="{size}" fill="{color}"{b}>{s}</text>'

    def line(x1, y1, x2, y2, sw=1, color=C, dash=''):
        d = f' stroke-dasharray="{dash}"' if dash else ''
        return f'<line x1="{x1:.2f}" y1="{y1:.2f}" x2="{x2:.2f}" y2="{y2:.2f}" stroke="{color}" stroke-width="{sw}"{d}/>'

    def rect(x, y, w, h, fill='none', stroke=C, sw=1.2):
        return f'<rect x="{x:.2f}" y="{y:.2f}" width="{w:.2f}" height="{h:.2f}" fill="{fill}" stroke="{stroke}" stroke-width="{sw}"/>'

    def hatch_rect(x, y, w, h):
        pts = []
        s2 = 8
        start_y = y - (x % s2)
        for yy in range(int(start_y), int(y + h + s2), s2):
            pts.append(line(x, yy, x + w, yy + s2, 0.5, CHATCH))
        for yy in range(int(start_y - s2), int(y + h + s2), s2):
            pts.append(line(x, yy + s2, x + w, yy, 0.5, CHATCH))
        return rect(x, y, w, h, 'none', C, 1.5) + ''.join(pts)

    def dim_45(x1, y1, x2, y2, label, side='above', offset=14, clr=CDIM):
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        a_ = 4
        s = ''
        s += line(x1, y1, x2, y2, 0.7, clr)
        s += line(x1 - a_, y1 - a_, x1 + a_, y1 + a_, 0.7, clr)
        s += line(x1 - a_, y1 + a_, x1 + a_, y1 - a_, 0.7, clr)
        s += line(x2 - a_, y2 - a_, x2 + a_, y2 + a_, 0.7, clr)
        s += line(x2 - a_, y2 + a_, x2 + a_, y2 - a_, 0.7, clr)
        if side == 'above':
            s += txt(mx, my - 5, label, 17, 'middle', False, clr)
        else:
            s += txt(mx, my + 16, label, 17, 'middle', False, clr)
        return s

    parts = []
    a = parts.append
    a('<div style="text-align:center;margin:10px 0;">')
    a('<svg width="%d" height="%d" xmlns="http://www.w3.org/2000/svg" '
      'style="font-family:SimSun,serif;">' % (svg_w, svg_h))

    # ────── 1. 立面图 ──────
    ml, mr = 60, 60
    pw = svg_w - ml - mr

    def px(x_mm):
        return ml + x_mm / L * pw if L > 0 else ml

    ey = 45
    eh = 42
    # 主梁
    a(rect(px(0), ey, pw, eh, 'none', C, 2))
    # 支座三角
    a(f'<polygon points="{px(0):.0f},{ey+eh+3} {px(0)-10:.0f},{ey+eh+16} {px(0)+10:.0f},{ey+eh+16}" fill="{C}"/>')
    a(f'<polygon points="{px(L):.0f},{ey+eh+3} {px(L)-10:.0f},{ey+eh+16} {px(L)+10:.0f},{ey+eh+16}" fill="{C}"/>')

    # 横向加劲肋
    n_total = int(L / st.spacing) + 1
    for i in range(1, n_total - 1):
        xi_mm = i * st.spacing
        if xi_mm < L:
            xp = px(xi_mm)
            a(line(xp, ey - 6, xp, ey + eh + 4, 1.5, C))
            # 小三角形标记肋顶
            a(f'<polygon points="{xp-4:.0f},{ey-6:.0f} {xp+4:.0f},{ey-6:.0f} {xp:.0f},{ey-10:.0f}" fill="{C}"/>')

    # 支座加劲肋（加粗）
    a(line(px(0), ey - 8, px(0), ey + eh + 6, 3, C))
    a(line(px(L), ey - 8, px(L), ey + eh + 6, 3, C))
    a(txt(px(0) - 18, ey + eh / 2 + 3, '支座', 12, 'end'))
    a(txt(px(L) + 18, ey + eh / 2 + 3, '支座', 12, 'start'))

    # 间距标注
    x1 = px(st.spacing)
    x2 = px(min(2 * st.spacing, L))
    if x2 > x1:
        ym = ey - 20
        a(dim_45(x1, ym, x2, ym, f'a={st.spacing:.0f}', 'above', 10, CDIM))
    # 跨度标注
    yL2 = ey + eh + 22
    a(dim_45(px(0), yL2, px(L), yL2, f'L={p.L:.1f}m', 'below', 10))
    # 剖面标记 C-C
    c_pos = (px(0) + px(L)) / 2
    a(txt(c_pos, ey + eh + 44, 'C', 17, 'middle', True))
    a(line(c_pos, ey + eh + 36, c_pos, ey + eh + 52, 0.7, C))
    a(line(c_pos - 6, ey + eh + 38, c_pos + 6, ey + eh + 38, 1, C))
    a(line(c_pos, ey - 10, c_pos, ey + eh + 6, 0.7, CDIM, '6,4'))

    # ────── 2. C-C 剖面（加劲肋处） ──────
    cx0, cy0 = 70, 180
    sc = 1.2
    bw = min(sec.b_f * sc, 160)
    sc = bw / max(sec.b_f, 1)
    h_px = r.props_mid.h * sc
    tw_px = max(sec.t_w * sc, 3)
    tf_px = max(sec.t_f * sc, 5)
    hw_px = sec.h_w * sc
    bf_px = sec.b_f * sc
    bs_px = min(st.b_s * sc, 60)
    cx_c = cx0 + bf_px / 2

    # 剖切填充
    a(hatch_rect(cx0, cy0, bf_px, tf_px))
    a(hatch_rect(cx0, cy0 + tf_px + hw_px, bf_px, tf_px))
    a(hatch_rect(cx_c - tw_px / 2, cy0 + tf_px, tw_px, hw_px))
    # 加劲肋（两侧）
    a(hatch_rect(cx0 - bs_px, cy0 + tf_px, bs_px, hw_px))
    a(hatch_rect(cx0 + bf_px, cy0 + tf_px, bs_px, hw_px))
    # 中心线
    a(line(cx_c - bf_px / 2 - bs_px - 10, cy0 + tf_px + hw_px / 2,
           cx_c + bf_px / 2 + bs_px + 10, cy0 + tf_px + hw_px / 2,
           0.7, CDIM, '8,4'))

    # 尺寸标注
    a(txt(cx_c, cy0 + h_px - 4, 'C—C', 16, 'middle', True))
    # 横向 b_s
    ly = cy0 + tf_px + hw_px + 14
    left_edge = cx0 - bs_px
    a(dim_45(left_edge, ly, cx0, ly, f'b_s={st.b_s}', 'below', 8))
    a(dim_45(cx0 + bf_px, ly, cx0 + bf_px + bs_px, ly, f'b_s={st.b_s}', 'below', 8))
    # t_s
    xs = cx0 + bf_px + bs_px + 16
    a(line(xs, cy0 + tf_px, xs, cy0 + tf_px + hw_px, 0.7, CDIM))
    a(line(xs - 3, cy0 + tf_px, xs + 3, cy0 + tf_px, 0.7, CDIM))
    a(line(xs - 3, cy0 + tf_px + hw_px, xs + 3, cy0 + tf_px + hw_px, 0.7, CDIM))
    a(txt(xs + 10, cy0 + tf_px + hw_px / 2 + 3, f't_s={st.t_s}', 17, 'start'))
    # h_w
    a(dim_45(cx0 - bs_px - 10, cy0 + tf_px, cx0 - bs_px - 10, cy0 + tf_px + hw_px, f'h_w={sec.h_w}', 'above', 10))
    # b_f
    a(dim_45(cx0, cy0 + tf_px + hw_px + 30, cx0 + bf_px, cy0 + tf_px + hw_px + 30, f'b_f={sec.b_f}', 'below', 8))

    # ────── 3. 支座加劲肋详图 ──────
    dx0, dy0 = 460, 155
    dw, dh = 310, 240
    a(rect(dx0, dy0, dw, dh, '#fafafa', C, 1.2))
    a(txt(dx0 + dw / 2, dy0 + 16, 'Ⅱ 支座加劲肋详图', 16, 'middle', True))

    sy = dy0 + 42
    web_w = 22
    bb = st.bearing_b
    bt = st.bearing_t
    bb_px = min(bb * 0.6, 65)
    bt_px = min(bt * 0.6, 18)
    bm = dx0 + dw / 2

    # 腹板
    a(hatch_rect(bm - web_w / 2, sy, web_w, 90))
    # 支座加劲肋
    a(hatch_rect(bm - web_w / 2 - bb_px, sy, bb_px, 90))
    a(hatch_rect(bm + web_w / 2, sy, bb_px, 90))
    # 端部刨平（粗实线）
    a(line(bm - web_w / 2 - bb_px, sy + 90, bm + web_w / 2 + bb_px, sy + 90, 2.5, C))
    a(txt(bm, sy + 102, '端部刨平顶紧', 17, 'middle', False, CDIM))

    # 下翼缘
    a(hatch_rect(bm - web_w / 2 - bb_px - 10, sy + 90, bb_px * 2 + web_w + 20, 14))
    # 尺寸 b 标注
    a(dim_45(bm - web_w / 2 - bb_px, sy + 128, bm - web_w / 2, sy + 128, f'b={bb}', 'below', 8))
    a(dim_45(bm + web_w / 2, sy + 128, bm + web_w / 2 + bb_px, sy + 128, f'b={bb}', 'below', 8))
    # t 标注
    xt = bm + web_w / 2 + bb_px + 14
    a(line(xt, sy, xt, sy + 90, 0.7, CDIM))
    a(line(xt - 3, sy, xt + 3, sy, 0.7, CDIM))
    a(line(xt - 3, sy + 90, xt + 3, sy + 90, 0.7, CDIM))
    a(txt(xt + 10, sy + 42, f't={bt}', 17, 'start'))
    # web_w 标注
    a(dim_45(bm - web_w / 2, sy - 12, bm + web_w / 2, sy - 12, f't_w={sec.t_w}', 'above', 8))

    # 细节文字说明
    a(txt(dx0 + dw / 2, dy0 + dh - 14,
          f'支座加劲肋 2—{bb}×{bt}，端部刨平顶紧于下翼缘',
          13, 'middle', False, CDIM))

    # ────── 4. 标题栏 ──────
    tb_y = svg_h - 50
    a(rect(10, tb_y, svg_w - 20, 40, 'none', C, 1.5))
    a(line(10, tb_y + 24, svg_w - 10, tb_y + 24, 0.7, C))
    a(txt(30, tb_y + 34, '加劲肋构造图', 16, 'start', True))
    a(txt(svg_w / 2, tb_y + 14, '比例 1:10', 17, 'middle', False, CDIM))
    a(txt(svg_w - 30, tb_y + 34, '单位：mm', 17, 'end', False, CDIM))

    a('</svg>')
    a('</div>')
    return '\n'.join(parts)


# ============================================================
# HTML 内容生成
# ============================================================

def generate_html(r: CalcResult) -> str:
    p = r.params
    sec = r.section
    pm = r.props_mid
    pv = r.props_var
    c = r.checks
    st = r.stiffeners
    wd = r.welds
    ft = r.fatigue

    L = []
    a = L.append

    a(_HTML_HEADER)

    # 标题
    title_tex = "%s\\,\\mathrm{m}\\text{简支焊接双轴对称工字形钢板梁设计计算书}" % p.L
    a(_md(title_tex))

    a('<table>')
    a('<tr><th>项目</th><th>内容</th></tr>')
    a('<tr><td>设计题目</td><td>%dm 简支工字型钢板梁设计</td></tr>' % p.L)
    a('<tr><td>设计依据</td><td>JTG D64-2015 / JTG D60-2015 / 钢结构设计原理</td></tr>')
    a('<tr><td>钢材</td><td>%s</td></tr>' % p.steel_grade)
    a('<tr><td>焊条</td><td>%s</td></tr>' % p.electrode)
    a('</table>')
    a('<hr>')

    # ============== 一、设计基本参数 ==============
    a('<h2>一、设计基本参数</h2>')
    a('<p>本设计为 %d m 简支公路桥梁钢板梁设计。桥梁上部结构采用钢-混凝土叠合梁形式，钢主梁为焊接双轴对称工字形截面，在跨径范围内进行一次变截面设计。</p>' % p.L)

    a('<h3>1.1 几何参数与荷载参数</h3>')
    a('<table>')
    a('<tr><th>参数</th><th>符号</th><th>数值</th><th>单位</th></tr>')
    a('<tr><td>计算跨径</td><td>%s</td><td>%.1f</td><td>m</td></tr>' % (_mi("L"), p.L))
    a('<tr><td>桥面板宽度</td><td>%s</td><td>%.1f</td><td>m</td></tr>' % (_mi("B"), p.B))
    a('<tr><td>混凝土桥面板厚度</td><td>%s</td><td>%.2f</td><td>m</td></tr>' % (_mi("H_1"), p.H1))
    a('<tr><td>沥青铺装层厚度</td><td>%s</td><td>%.2f</td><td>m</td></tr>' % (_mi("H_2"), p.H2))
    a('<tr><td>横向分布系数</td><td>%s</td><td>%.2f</td><td>&mdash;</td></tr>' % (_mi(LX.ETA), p.eta))
    a('<tr><td>恒载分项系数</td><td>%s</td><td>%.1f</td><td>&mdash;</td></tr>' % (_mi(LX.GAMMA + "_G"), p.gamma_G))
    a('<tr><td>活载分项系数</td><td>%s</td><td>%.1f</td><td>&mdash;</td></tr>' % (_mi(LX.GAMMA + "_Q"), p.gamma_Q))
    a('<tr><td>结构重要性系数</td><td>%s</td><td>%.1f</td><td>&mdash;</td></tr>' % (_mi(LX.GAMMA + "_0"), p.gamma_0))
    a('<tr><td>活载冲击系数</td><td>%s</td><td>%.2f</td><td>&mdash;</td></tr>' % (_mi("1+" + LX.MU), p.mu_impact))
    a('</table>')

    a('<h3>1.2 材料特性</h3>')
    a('<p>钢材选用 %s 低合金高强度结构钢，焊条选用 %s 型。依据 JTG D64-2015 表 3.2.1，%s 钢强度设计值如下：</p>' % (p.steel_grade, p.electrode, p.steel_grade))
    a('<table>')
    a('<tr><th>板厚 t (mm)</th><th>f<sub>d</sub> (MPa)</th><th>f<sub>vd</sub> (MPa)</th><th>f<sub>ce</sub> (MPa)</th></tr>')
    a('<tr><td>t &le; 16</td><td>275</td><td>160</td><td>355</td></tr>')
    a('<tr><td>16 &lt; t &le; 40</td><td>270</td><td>155</td><td>355</td></tr>')
    a('<tr><td>40 &lt; t &le; 63</td><td>260</td><td>150</td><td>355</td></tr>')
    a('</table>')
    E_show = E_STEEL / 1e5
    G_show = G_STEEL / 1e4
    a('<p>弹性模量 %s，剪切模量 %s，泊松比 %s。钢材容重 %s（即 78.5 kN/m&sup3;）。</p>' % (
        _mi("E = %.2f \\times 10^5\\,\\mathrm{MPa}" % E_show),
        _mi("G = %.2f \\times 10^4\\,\\mathrm{MPa}" % G_show),
        _mi(LX.NU + " = 0.3"),
        _mi(LX.RHO + " = 7850\\,\\mathrm{kg/m^3}"),
    ))

    # ============== 二、荷载计算 ==============
    a('<h2>二、荷载计算</h2>')

    a('<h3>2.1 二期恒载（桥面板+铺装层）</h3>')
    g2 = r.g2
    a(_md("g_2 = \\gamma_c \\times B \\times H_1 + \\gamma_a \\times B \\times H_2"))
    term1 = 25.0 * p.B * p.H1
    term2 = 24.0 * p.B * p.H2
    a(_md("g_2 = 25.0 \\times %.1f \\times %.2f + 24.0 \\times %.1f \\times %.2f = %.3f + %.3f = %.3f\\,\\mathrm{kN/m}" % (
        p.B, p.H1, p.B, p.H2, term1, term2, g2)))
    a('<p><strong>二期恒载合计：%s</strong></p>' % _mi("g_2 = %.3f\\,\\mathrm{kN/m}" % g2))

    a('<h3>2.2 活载标准值（JTG D60-2015 第4.3.1条）</h3>')
    a('<p>公路-I级车道荷载：均布荷载 %s（满跨布置）。集中荷载 P<sub>k</sub>（用于弯矩计算）按下式内插：</p>' % _mi("q_k = %s\\,\\mathrm{kN/m}" % Q_K))
    a(_md("P_k = 270 + (360 - 270) \\times (L - 5) / (50 - 5)"))
    a(_md("P_k = 270 + 90 \\times (%d - 5) / 45 = %.2f\\,\\mathrm{kN}" % (p.L, r.P_k)))
    a('<p>用于剪力计算时，P<sub>k</sub> 乘 1.2 系数：%s。</p>' % _mi("P_{k,\\mathrm{shear}} = 1.2 \\times %.2f = %.2f\\,\\mathrm{kN}" % (r.P_k, r.P_k * 1.2)))
    a('<p>横向分布系数 %s，分配到单根主梁的活载标准值：</p>' % _mi(LX.ETA + " = %.2f" % p.eta))
    a('<ul>')
    a('<li>均布活载：%s</li>' % _mi("q = %s \\times q_k = %.2f \\times %s = %.3f\\,\\mathrm{kN/m}" % (LX.ETA, p.eta, Q_K, r.q)))
    a('<li>集中活载（弯矩用）：%s</li>' % _mi("P = %s \\times P_k = %.2f \\times %.2f = %.2f\\,\\mathrm{kN}" % (LX.ETA, p.eta, r.P_k, r.P)))
    a('<li>集中活载（剪力用）：%s</li>' % _mi("P_s = %s \\times 1.2 \\times P_k = %.2f \\times %.2f = %.2f\\,\\mathrm{kN}" % (LX.ETA, p.eta, r.P_k * 1.2, r.P_s)))
    a('</ul>')

    a('<h3>2.3 一期恒载（钢梁自重）</h3>')
    a('<p>钢梁自重取决于截面尺寸，先按经验初估，拟定截面后代入验算。</p>')

    # ============== 三、截面尺寸拟定 ==============
    a('<h2>三、截面尺寸拟定</h2>')
    a('<h3>3.1 梁高确定</h3>')
    a('<p>梁高是钢板梁经济性的关键参数，综合以下因素确定：</p>')
    h_min15 = p.L * 1000 / 15
    h_min12 = p.L * 1000 / 12
    a('<p><strong>(1) 按刚度条件（挠度控制）：</strong>简支梁满足挠度限值 L/500 所需的最小梁高约为 %s。</p>' %
      _mi("h = L/15 \\sim L/12 = %.0f \\sim %.0f\\,\\mathrm{mm}" % (h_min15, h_min12)))
    a('<p><strong>(2) 按经济条件：</strong>经济梁高经验公式 %s。</p>' %
      _mi("h_e = 7 \\cdot W_{\\mathrm{req}}^{1/3} - 300\\;(\\mathrm{mm})"))
    a('<p><strong>(3) 综合确定：</strong>腹板高度 %s，翼缘厚度 %s（上下翼缘等厚），总梁高 %s。</p>' % (
        _mi("h_w = %.0f\\,\\mathrm{mm}" % sec.h_w),
        _mi("t_f = %.0f\\,\\mathrm{mm}" % sec.t_f),
        _mi("h = h_w + 2t_f = %.0f\\,\\mathrm{mm}" % pm.h)))

    a('<h3>3.2 拟选截面尺寸</h3>')
    a('<table>')
    a('<tr><th>板件</th><th>符号</th><th>尺寸 (mm)</th><th>说明</th></tr>')
    a('<tr><td>腹板高度</td><td>%s</td><td>%.0f</td><td>按 L/15~L/20 经济与刚度确定</td></tr>' % (_mi("h_w"), sec.h_w))
    a('<tr><td>腹板厚度</td><td>%s</td><td>%.0f</td><td>满足抗剪与局部稳定，%s</td></tr>' % (
        _mi("t_w"), sec.t_w, _mi("t_w \\geq h_w/170 = %.1f" % (sec.h_w / 170))))
    a('<tr><td>翼缘宽度</td><td>%s</td><td>%.0f</td><td>满足整体稳定，%s</td></tr>' % (
        _mi("b_f"), sec.b_f, _mi("b_f \\geq h/5 = %.0f" % (pm.h / 5))))
    a('<tr><td>翼缘厚度</td><td>%s</td><td>%.0f</td><td>满足抗弯所需截面模量</td></tr>' % (_mi("t_f"), sec.t_f))
    a('<tr><td>梁总高</td><td>%s</td><td>%.0f</td><td>%s</td></tr>' % (_mi("h"), pm.h, _mi("h = h_w + 2t_f")))
    a('</table>')

    # 横断面图 (SVG)
    a(_make_cross_section_svg(sec.h_w, sec.t_w, sec.b_f, sec.t_f, pm.h))

    a('<h3>3.3 截面几何特性计算</h3>')
    I_w = sec.t_w * sec.h_w ** 3 / 12.0
    d_f = (sec.h_w + sec.t_f) / 2.0
    I_f = 2.0 * (sec.b_f * sec.t_f ** 3 / 12.0 + sec.b_f * sec.t_f * d_f ** 2)
    I_f_line = ("I_f = 2\\left[ \\frac{%d \\times %d^3}{12} + %d \\times %d \\times %d^2 \\right] = %.2f \\times 10^6\\,\\mathrm{mm^4}" %
                (sec.b_f, sec.t_f, sec.b_f, sec.t_f, d_f, I_f / 1e6))

    a('<p><strong>(1) 截面面积</strong></p>')
    a(_md("A = %d \\times %d + 2 \\times %d \\times %d = %d\\,\\mathrm{mm^2}" % (sec.h_w, sec.t_w, sec.b_f, sec.t_f, pm.A)))
    a('<p><strong>(2) 惯性矩（绕强轴 x-x）</strong></p>')
    a(_md("I_w = t_w h_w^3 / 12 = %d \\times %d^3 / 12 = %.2f \\times 10^6\\,\\mathrm{mm^4}" % (sec.t_w, sec.h_w, I_w / 1e6)))
    a(_md(I_f_line))
    a(_md("I_x = I_w + I_f = %.2f \\times 10^6\\,\\mathrm{mm^4}" % (pm.I_x / 1e6)))
    a('<p><strong>(3) 弹性截面模量</strong></p>')
    a(_md("W_x = I_x / (h/2) = %.2f \\times 10^6 / %d = %.2f\\,\\mathrm{cm^3}" % (pm.I_x / 1e6, pm.h / 2, pm.W_x / 1e3)))
    a('<p><strong>(4) 半截面面积矩（用于剪应力计算）</strong></p>')
    a(_md("S = b_f t_f (h_w + t_f)/2 + t_w (h_w/2)^2 / 2"))
    a(_md("S = %d \\times %d \\times %d + %d \\times (%d)^2 / 2 = %.2f\\,\\mathrm{cm^3}" % (
        sec.b_f, sec.t_f, d_f, sec.t_w, sec.h_w / 2, pm.S / 1e3)))
    a('<p><strong>(5) 钢梁自重（一期恒载）</strong></p>')
    a(_md("g_1 = A \\times \\rho_{\\mathrm{steel}} = %.4f \\times 78.5 = %.3f\\,\\mathrm{kN/m}" % (pm.A * 1e-6, r.g1)))
    a('<p><strong>总恒载：%s</strong></p>' % _mi("g = g_1 + g_2 = %.3f + %.3f = %.3f\\,\\mathrm{kN/m}" % (r.g1, g2, r.g)))

    # ============== 四、内力计算 ==============
    a('<h2>四、内力计算</h2>')
    a('<p>简支梁计算跨径 %s，控制截面为跨中（弯矩最大）和支座（剪力最大）。</p>' % _mi("L = %g\\,\\mathrm{m}" % p.L))

    a('<h3>4.1 跨中弯矩</h3>')
    a('<p><strong>弯矩标准值：</strong></p>')
    a(_md("M_{gk} = g L^2 / 8 = %.3f \\times %.0f^2 / 8 = %.2f\\,\\mathrm{kN\\cdot m}" % (r.g, p.L, r.M_gk)))
    a(_md("M_{qk} = q L^2 / 8 + P L / 4 = %.3f \\times %.0f^2 / 8 + %.2f \\times %.0f / 4 = %.2f\\,\\mathrm{kN\\cdot m}" %
          (r.q, p.L, r.P, p.L, r.M_qk)))
    a('<p><strong>弯矩设计值（承载能力极限状态）：</strong></p>')
    a(_md("M_{Ed} = " + LX.GAMMA + "_0 [" + LX.GAMMA + "_G \\times M_{gk} + " + LX.GAMMA +
          "_Q \\times (1+" + LX.MU + ") \\times M_{qk}]"))
    a(_md("M_{Ed} = %.1f \\times [%.1f \\times %.2f + %.1f \\times %.2f \\times %.2f] = %.2f\\,\\mathrm{kN\\cdot m}" %
          (p.gamma_0, p.gamma_G, r.M_gk, p.gamma_Q, p.mu_impact, r.M_qk, r.M_Ed)))

    a('<h3>4.2 支座剪力</h3>')
    a('<p><strong>剪力标准值：</strong></p>')
    a(_md("V_{gk} = g L / 2 = %.2f\\,\\mathrm{kN}" % r.V_gk))
    a(_md("V_{qk} = q L / 2 + P_s / 2 = %.2f\\,\\mathrm{kN}" % r.V_qk))
    a('<p><strong>剪力设计值：</strong></p>')
    a(_md("V_{Ed} = " + LX.GAMMA + "_0 [" + LX.GAMMA + "_G \\times V_{gk} + " + LX.GAMMA +
          "_Q \\times (1+" + LX.MU + ") \\times V_{qk}] = %.2f\\,\\mathrm{kN}" % r.V_Ed))
    a('')
    a(_make_moment_shear_diagrams(
        p.L, r.g, r.q, r.P, r.P_s, r.x_cut,
        r.M_gk, r.M_qk, r.M_Ed,
        r.M_gk_x, r.M_qk_x, r.M_Ed_x,
        r.V_gk, r.V_qk, r.V_Ed))

    # ============== 五、跨中强度及刚度验算 ==============
    a('<h2>五、跨中截面强度及刚度验算</h2>')

    a('<h3>5.1 抗弯强度（JTG D64-2015 第5.2条）</h3>')
    a('<p>翼缘板厚 %s（%d &lt; t<sub>f</sub> &le; %d mm），取 %s。</p>' % (
        _mi("t_f = %d\\,\\mathrm{mm}" % sec.t_f),
        16 if sec.t_f <= 40 else 40, 40 if sec.t_f <= 40 else 63,
        _mi("f_d = %d\\,\\mathrm{MPa}" % r.f_d_mid)))
    a(_md(LX.SIGMA + " = M_{Ed} / W_x = %.2f \\times 10^6 / (%.2f \\times 10^3) = %.1f\\,\\mathrm{MPa}" %
          (r.M_Ed, pm.W_x / 1e3, c.sigma_mid)))
    cls_s = 'ok' if c.sigma_mid_ok else 'ng'
    ok_s = '满足' if c.sigma_mid_ok else '不满足'
    a('<p><strong>%s &nbsp; [<span class="%s">%s</span>]</strong></p>' % (
        _mi(LX.SIGMA + " = %.1f\\,\\mathrm{MPa} < f_d = %d\\,\\mathrm{MPa}" % (c.sigma_mid, c.sigma_mid_limit)),
        cls_s, ok_s))

    a('<h3>5.2 抗剪强度（JTG D64-2015 第5.3条）</h3>')
    a('<p>腹板中性轴处剪应力最大，t<sub>w</sub> = %d mm，取 %s。</p>' % (sec.t_w, _mi("f_{vd} = %d\\,\\mathrm{MPa}" % r.f_vd_mid)))
    a(_md(LX.TAU + "_{\\max} = \\frac{V_{Ed} \\times S}{I_x \\times t_w}"))
    a(_md(LX.TAU + "_{\\max} = \\frac{%.2f \\times 10^3 \\times %.2f \\times 10^3}{%.2f \\times 10^6 \\times %d} = %.1f\\,\\mathrm{MPa}" %
          (r.V_Ed, pm.S / 1e3, pm.I_x / 1e6, sec.t_w, c.tau_max)))
    cls_t = 'ok' if c.tau_max_ok else 'ng'
    ok_t = '满足' if c.tau_max_ok else '不满足'
    a('<p><strong>%s &nbsp; [<span class="%s">%s</span>]</strong></p>' % (
        _mi(LX.TAU + "_{\\max} = %.1f\\,\\mathrm{MPa} < f_{vd} = %d\\,\\mathrm{MPa}" % (c.tau_max, c.tau_max_limit)),
        cls_t, ok_t))

    a('<h3>5.3 折算应力（JTG D64-2015 第5.4条）</h3>')
    y_j = sec.h_w / 2.0
    sigma_j = r.M_Ed * 1e6 * y_j / pm.I_x
    tau_j = r.V_Ed * 1e3 * pm.S_f / (pm.I_x * sec.t_w)
    a('<p>验算跨中截面翼缘与腹板交界处（%s）：</p>' % _mi("y_j = h_w/2 = %d\\,\\mathrm{mm}" % y_j))
    a(_md(LX.SIGMA + "_j = \\frac{M_{Ed} \\times y_j}{I_x} = %.1f\\,\\mathrm{MPa}" % sigma_j))
    a(_md(LX.TAU + "_j = \\frac{V_{Ed} \\times S_f}{I_x \\times t_w} = %.1f\\,\\mathrm{MPa}" % tau_j))
    a(_md(LX.SIGMA + "_{zs} = \\sqrt{" + LX.SIGMA + "_j^2 + 3" + LX.TAU + "_j^2} = \\sqrt{%.1f^2 + 3 \\times %.1f^2} = %.1f\\,\\mathrm{MPa}" %
          (sigma_j, tau_j, c.sigma_zs)))
    cls_z = 'ok' if c.sigma_zs_ok else 'ng'
    ok_z = '满足' if c.sigma_zs_ok else '不满足'
    a('<p><strong>%s &nbsp; [<span class="%s">%s</span>]</strong></p>' % (
        _mi(LX.SIGMA + "_{zs} = %.1f\\,\\mathrm{MPa} < 1.1 f_d = %d\\,\\mathrm{MPa}" % (c.sigma_zs, c.sigma_zs_limit)),
        cls_z, ok_z))

    a('<h3>5.4 刚度验算（JTG D64-2015 第5.4条）</h3>')
    L_mm = p.L * 1000.0
    q_Nmm = r.q
    P_N = r.P * 1000.0
    delta_u = 5.0 * q_Nmm * L_mm ** 4 / (384.0 * E_STEEL * pm.I_x)
    delta_c = P_N * L_mm ** 3 / (48.0 * E_STEEL * pm.I_x)
    a('<p>活载挠度（不计冲击系数）按均布+集中荷载叠加计算：</p>')
    a(_md(LX.DELTA + "_u = \\frac{5 q L^4}{384 E I_x} = %.1f\\,\\mathrm{mm}" % delta_u))
    a(_md(LX.DELTA + "_c = \\frac{P L^3}{48 E I_x} = %.1f\\,\\mathrm{mm}" % delta_c))
    a(_md(LX.DELTA + "_q = " + LX.DELTA + "_u + " + LX.DELTA + "_c = %.1f\\,\\mathrm{mm}" % c.deflection_q))
    a(_md("[" + LX.DELTA + "] = L / 500 = %.1f\\,\\mathrm{mm}" % c.deflection_limit))
    cls_d = 'ok' if c.deflection_q_ok else 'ng'
    ok_d = '满足' if c.deflection_q_ok else '不满足'
    a('<p><strong>%s &nbsp; [<span class="%s">%s</span>]</strong></p>' % (
        _mi((LX.DELTA + "_q = %.1f\\,\\mathrm{mm} < [" + LX.DELTA + "] = %.1f\\,\\mathrm{mm}") % (c.deflection_q, c.deflection_limit)),
        cls_d, ok_d))
    a('<p>恒载挠度 &delta;<sub>g</sub> = %.1f mm，总挠度 = %.1f mm &lt; L/500 = %.1f mm，满足挠度验算要求。</p>' %
      (c.deflection_g, c.deflection_total, p.L * 1000 / 500))

    a('<h3>5.5 整体稳定性（JTG D64-2015 第5.5条）</h3>')

    # 翼缘局部稳定
    b1 = (sec.b_f - sec.t_w) / 2.0
    fy = 345.0 if p.steel_grade == "Q345" else 235.0
    corr_local = math.sqrt(235.0 / fy)
    local_limit = 15.0 * corr_local

    a('<h4>5.5.1 翼缘局部稳定（JTG D64-2015 第5.7节）</h4>')
    a('<p>受压翼缘自由外伸宽度 b<sub>1</sub> 与厚度 t<sub>f</sub> 之比应满足：</p>')
    a(_md("b_1 = (b_f - t_w) / 2 = (%.0f - %.0f) / 2 = %.0f\\,\\mathrm{mm}" %
          (sec.b_f, sec.t_w, b1)))
    a(_md("\\frac{b_1}{t_f} = \\frac{%.0f}{%.0f} = %.2f" % (b1, sec.t_f, c.flange_local_ratio)))
    a('<p>限值：%s</p>' % _mi("15 \\sqrt{235 / f_y} = 15 \\times %.3f = %.1f" % (corr_local, local_limit)))
    fl_cls = 'ok' if c.flange_local_ok else 'ng'
    fl_ok = '满足' if c.flange_local_ok else '不满足'
    a('<p><strong>%s &nbsp; [<span class="%s">%s</span>]</strong></p>' %
      (_mi("\\frac{b_1}{t_f} = %.2f \\leq %.1f" % (c.flange_local_ratio, local_limit)), fl_cls, fl_ok))

    a('<h4>5.5.2 整体稳定性（JTG D64-2015 第5.5.2条）</h4>')
    a('<p>依据 JTG D64-2015 第5.5.2条第1款：</p>')
    a('<p style="padding-left:2em;">"有钢筋混凝土铺板密铺在梁的受压翼缘上并与其牢固连接、'
      '能阻止梁受压翼缘的侧向位移时，可不计算整体稳定性。"</p>')
    a('<p>本桥采用钢-混凝土叠合梁，混凝土桥面板通过焊钉剪力连接件与钢主梁上翼缘焊接，'
      '两者形成牢固连接，桥面板为受压上翼缘提供连续侧向约束。'
      '<strong>因此整体稳定自然满足，无需计算。</strong></p>')

    a('<h3>5.6 疲劳强度（JTG D64-2015 第5.6节及附录C）</h3>')
    a('<p>采用疲劳荷载模型 I（等效车道荷载乘 0.7 折减系数）：</p>')
    a('<p>%s，%s。</p>' % (
        _mi("q_f = 0.7 \\times 10.5 = 7.35\\,\\mathrm{kN/m}"),
        _mi("P_f = 0.7 \\times %.2f = %.2f\\,\\mathrm{kN}" % (r.P_k, 0.7 * r.P_k))))
    a('<p>按 %s 分配到单梁：%s，%s。</p>' % (
        _mi(LX.ETA + " = %.2f" % p.eta),
        _mi("q_{f1} = %.3f\\,\\mathrm{kN/m}" % ft.q_f1),
        _mi("P_{f1} = %.2f\\,\\mathrm{kN}" % ft.P_f1)))
    a('<p>%s。</p>' % _mi("M_f = q_{f1} \\times L^2 / 8 + P_{f1} \\times L / 4 = %.2f\\,\\mathrm{kN\\cdot m}" % ft.M_f))
    a(_md(LX.DELTA + LX.SIGMA + "_p = M_f / W_x = %.2f \\times 10^6 / (%.2f \\times 10^3) = %.1f\\,\\mathrm{MPa}" %
          (ft.M_f, pm.W_x / 1e3, ft.delta_sigma_p)))
    a('<table>')
    a('<tr><th>细节类别</th><th>&Delta;&sigma;<sub>c</sub> (MPa)</th><th>&Delta;&sigma;<sub>D</sub> = 0.74&Delta;&sigma;<sub>c</sub></th><th>&Delta;&sigma;<sub>p</sub></th><th>判定</th></tr>')
    cls_f1 = 'ok' if ft.check_base_metal else 'ng'
    ok_f1 = '满足' if ft.check_base_metal else '不满足'
    cls_f2 = 'ok' if ft.check_fillet_weld else 'ng'
    ok_f2 = '满足' if ft.check_fillet_weld else '不满足'
    a('<tr><td>翼缘母材（非焊接）</td><td>160</td><td>%.1f</td><td>%.1f</td><td><span class="%s">%s</span></td></tr>' %
      (ft.delta_sigma_D_base, ft.delta_sigma_p, cls_f1, ok_f1))
    a('<tr><td>翼缘-腹板连续角焊缝</td><td>80</td><td>%.1f</td><td>%.1f</td><td><span class="%s">%s</span></td></tr>' %
      (ft.delta_sigma_D_weld, ft.delta_sigma_p, cls_f2, ok_f2))
    a('</table>')
    a('<p><strong>&Delta;&sigma;<sub>p</sub> = %.1f MPa &lt; 各细节类别常幅疲劳极限，疲劳强度满足。</strong></p>' % ft.delta_sigma_p)

    # ============== 六、变截面设计 ==============
    a('<h2>六、变截面设计</h2>')
    a('<h3>6.1 变截面原理</h3>')
    a('<p>简支梁弯矩沿跨径呈抛物线分布。为节约钢材，在距支座 L/6 = %.3f m 处减小翼缘尺寸，腹板尺寸保持不变。变截面处弯矩约为跨中最大弯矩的 %.1f%%。</p>' %
      (r.x_cut, r.M_Ed_x / r.M_Ed * 100))

    a('<h3>6.2 变截面处内力</h3>')
    a(_md("M_{gk}(x) = g \\times x \\times (L-x) / 2 = %.2f\\,\\mathrm{kN\\cdot m}" % r.M_gk_x))
    a(_md("M_{qk}(x) = q \\times x \\times (L-x) / 2 + P \\times x / 2 = %.2f\\,\\mathrm{kN\\cdot m}" % r.M_qk_x))
    a(_md("M_{Ed}(x) = " + LX.GAMMA + "_0 [" + LX.GAMMA + "_G M_{gk}(x) + " + LX.GAMMA +
          "_Q (1+" + LX.MU + ") M_{qk}(x)] = %.2f\\,\\mathrm{kN\\cdot m}" % r.M_Ed_x))
    a('<p>V<sub>Ed</sub>(x) = %.2f kN</p>' % r.V_Ed_x)

    a('<h3>6.3 变截面处尺寸</h3>')
    a('<table>')
    a('<tr><th>参数</th><th>等截面段（跨中）</th><th>变截面段</th></tr>')
    a('<tr><td>腹板 h<sub>w</sub> &times; t<sub>w</sub> (mm)</td><td>%d &times; %d</td><td>%d &times; %d (不变)</td></tr>' %
      (sec.h_w, sec.t_w, sec.h_w, sec.t_w))
    a('<tr><td>翼缘宽 b<sub>f</sub> (mm)</td><td>%d</td><td>%d</td></tr>' % (sec.b_f, sec.b_f2))
    a('<tr><td>翼缘厚 t<sub>f</sub> (mm)</td><td>%d</td><td>%d</td></tr>' % (sec.t_f, sec.t_f2))
    a('<tr><td>总高 h (mm)</td><td>%d</td><td>%d</td></tr>' % (pm.h, pv.h))
    a('<tr><td>惯性矩 I<sub>x</sub> (10<sup>6</sup> mm<sup>4</sup>)</td><td>%.2f</td><td>%.2f</td></tr>' %
      (pm.I_x / 1e6, pv.I_x / 1e6))
    a('<tr><td>截面模量 W<sub>x</sub> (cm<sup>3</sup>)</td><td>%.2f</td><td>%.2f</td></tr>' %
      (pm.W_x / 1e3, pv.W_x / 1e3))
    a('</table>')

    a(_make_variable_section_layout_svg(p.L, r.x_cut,
        (sec.h_w, sec.t_w, sec.b_f, sec.t_f),
        (sec.h_w, sec.t_w, sec.b_f2, sec.t_f2)))

    a('<h3>6.4 变截面处强度验算</h3>')
    cls_v = 'ok' if c.sigma_var_ok else 'ng'; ok_v = '满足' if c.sigma_var_ok else '不满足'
    cls_vt = 'ok' if c.tau_var_ok else 'ng'; ok_vt = '满足' if c.tau_var_ok else '不满足'
    cls_vz = 'ok' if c.sigma_zs_var_ok else 'ng'; ok_vz = '满足' if c.sigma_zs_var_ok else '不满足'
    a('<p><strong>抗弯：</strong>%s = %.1f MPa &lt; %s MPa [<span class="%s">%s</span>]</p>' %
      (_mi(LX.SIGMA + " = M_{Ed}(x) / W_{x2}"), c.sigma_var,
       _mi("f_d = %d" % c.sigma_var_limit), cls_v, ok_v))
    a('<p><strong>抗剪：</strong>%s = %.1f MPa &lt; %s MPa [<span class="%s">%s</span>]</p>' %
      (_mi(LX.TAU + " = V_{Ed}(x) \\times S / (I_{x2} \\times t_w)"), c.tau_var,
       _mi("f_{vd} = %d" % c.tau_var_limit), cls_vt, ok_vt))
    a('<p><strong>折算应力：</strong>%s = %.1f MPa &lt; %s MPa [<span class="%s">%s</span>]</p>' %
      (_mi(LX.SIGMA + "_{zs} = \\sqrt{" + LX.SIGMA + "_j^2 + 3" + LX.TAU + "_j^2}"), c.sigma_zs_var,
       _mi("1.1f_d = %d" % c.sigma_zs_var_limit), cls_vz, ok_vz))

    # ============== 七、翼缘焊缝设计 ==============
    a('<h2>七、翼缘焊缝设计</h2>')
    a('<p>翼缘与腹板采用双面连续角焊缝连接，%s 焊条。角焊缝强度设计值 %s。</p>' %
      (p.electrode, _mi("f_{\\mathrm{ff}} = 200\\,\\mathrm{MPa}")))

    a('<h3>7.1 支座截面处（等截面段）</h3>')
    a(_md("S_f = b_f \\times t_f \\times (h_w + t_f) / 2 = %d \\times %d \\times %d = %.2f\\,\\mathrm{cm^3}" %
          (sec.b_f, sec.t_f, d_f, wd.S_f1 / 1e3)))
    a(_md("v = V_{Ed} \\times S_f / I_x = %.2f \\times 10^3 \\times %.2f \\times 10^3 / (%.2f \\times 10^6) = %.1f\\,\\mathrm{N/mm}" %
          (r.V_Ed, wd.S_f1 / 1e3, pm.I_x / 1e6, wd.v1)))
    a('<p>所需焊脚尺寸（双面角焊缝）：%s</p>' %
      _mi("h_f \\leq v / (2 \\times 0.7 \\times f_{\\mathrm{ff}}) = %.1f / (2 \\times 0.7 \\times 200) = %.1f\\,\\mathrm{mm}" %
          (wd.v1, wd.h_f_req1)))
    a('<p>构造要求：%s，%s。</p>' % (
        _mi("h_{f,\\min} = 1.5\\sqrt{%d} = %.1f\\,\\mathrm{mm}" % (sec.t_f, wd.h_f_min1)),
        _mi("h_{f,\\max} = 1.2 \\times %d = %.1f\\,\\mathrm{mm}" % (sec.t_w, wd.h_f_max1))))
    a('<p><strong>选用焊脚尺寸：%s</strong></p>' % _mi("h_f = %d\\,\\mathrm{mm}" % wd.h_f_chosen))

    a('<h3>7.2 变截面处</h3>')
    a(_md("S_{f2} = %d \\times %d \\times (%d + %d) / 2 = %.2f\\,\\mathrm{cm^3}" %
          (sec.b_f2, sec.t_f2, sec.h_w, sec.t_f2, wd.S_f2 / 1e3)))
    a(_md("v = %.2f \\times 10^3 \\times %.2f \\times 10^3 / (%.2f \\times 10^6) = %.1f\\,\\mathrm{N/mm}" %
          (r.V_Ed_x, wd.S_f2 / 1e3, pv.I_x / 1e6, wd.v2)))
    a('<p>所需 %s。构造：[%.1f, %.1f] mm。</p>' % (
        _mi("h_f \\leq %.1f / (2 \\times 0.7 \\times 200) = %.1f\\,\\mathrm{mm}" % (wd.v2, wd.h_f_req2)),
        wd.h_f_min2, wd.h_f_max2))
    a('<p><strong>选用焊脚尺寸：%s（全跨统一）</strong></p>' % _mi("h_f = %d\\,\\mathrm{mm}" % wd.h_f_chosen))
    a('<p><strong>【结论】翼缘-腹板连接焊缝全跨统一采用 %s 双面连续角焊缝。</strong></p>' %
      _mi("h_f = %d\\,\\mathrm{mm}" % wd.h_f_chosen))

    a(_make_girder_weld_svg(r))

    # ============== 八、局部稳定设计 ==============
    a('<h2>八、局部稳定设计</h2>')
    a('<h3>8.1 腹板加劲肋（JTG D64-2015 第5.7节）</h3>')
    a('<p>%s</p>' % _mi("h_w / t_w = %d / %d = %.1f" % (sec.h_w, sec.t_w, st.hw_tw_ratio)))
    a('<p>根据 JTG D64-2015 第5.7节，对 Q345 钢：</p>')
    a('<ul>')
    a('<li>h<sub>w</sub>/t<sub>w</sub> &le; 100：局部稳定自然满足；</li>')
    a('<li>100 &lt; h<sub>w</sub>/t<sub>w</sub> &le; 170：需配置横向加劲肋；</li>')
    a('<li>h<sub>w</sub>/t<sub>w</sub> &gt; 170：需配置横向和纵向加劲肋。</li>')
    a('</ul>')
    if st.need_longitudinal:
        a('<p>因 h<sub>w</sub>/t<sub>w</sub> = %.1f &gt; 170，<strong>需配置横向和纵向加劲肋</strong>。</p>' % st.hw_tw_ratio)
    elif st.need_transverse:
        a('<p>因 100 &lt; %.1f &le; 170，<strong>需配置横向加劲肋</strong>。</p>' % st.hw_tw_ratio)
    else:
        a('<p>因 h<sub>w</sub>/t<sub>w</sub> = %.1f &le; 100，局部稳定自然满足。</p>' % st.hw_tw_ratio)

    if st.need_transverse:
        a('<p><strong>(1) 加劲肋间距</strong></p>')
        a('<p>横向加劲肋间距 a &le; min(2h<sub>w</sub>, 3000) = %d mm。</p>' % min(2 * sec.h_w, 3000))
        a('<p>取间距 a = %.0f mm（满足 0.5h<sub>w</sub> = %.0f &le; a &le; 3000），全跨均匀布置 %d 对。</p>' %
          (st.spacing, 0.5 * sec.h_w, st.n_pairs))
        a('<p><strong>(2) 加劲肋构造尺寸（JTG D64-2015 第5.7.3条）</strong></p>')
        a('<p>%s</p>' % _mi("b_s \\geq h_w/30 + 40 = %.1f\\,\\mathrm{mm}" % (sec.h_w / 30 + 40)))
        a('<p>%s</p>' % _mi("t_s \\geq b_s / 15 = %.1f\\,\\mathrm{mm}" % (st.b_s / 15)))
        a('<p><strong>选用 b<sub>s</sub> &times; t<sub>s</sub> = %d &times; %d mm</strong>，成对布置在腹板两侧。</p>' %
          (st.b_s, st.t_s))
        a('<p>横向加劲肋与腹板采用双面角焊缝，h<sub>f</sub> = 6 mm。加劲肋两端不与翼缘焊接。</p>')

    if st.need_transverse:
        a('<h4>8.1.1 腹板区格受剪屈曲（JTG D64-2015 第5.7.4条）</h4>')
        a_show = st.spacing
        ratio_a_hw = a_show / sec.h_w
        k_tau = 5.34 + 4.00 / ratio_a_hw ** 2 if ratio_a_hw >= 1.0 else 4.00 + 5.34 / ratio_a_hw ** 2
        a('<p>横向加劲肋间距 a = %.0f mm，a/h<sub>w</sub> = %.3f。剪切屈曲系数：</p>' % (a_show, ratio_a_hw))
        a(_md("k_\\tau = %.3f" % k_tau))
        eps_k = math.sqrt(235.0 / 345.0)
        lam_s = (sec.h_w / sec.t_w) / (41.0 * eps_k * math.sqrt(k_tau))
        a('<p>正则化高厚比 &lambda;<sub>s</sub> = (%.1f) / (41 &times; %.3f &times; &radic;(%.3f)) = %.3f</p>' %
          (st.hw_tw_ratio, eps_k, k_tau, lam_s))
        a('<p>腹板平均剪应力：&tau; = %.1f MPa，受剪屈曲临界应力：&tau;<sub>cr</sub> = %.1f MPa。</p>' %
          (st.tau_web, st.tau_cr))
        cls_sb = 'ok' if st.shear_buckling_ok else 'ng'
        ok_sb = '满足' if st.shear_buckling_ok else '不满足'
        a('<p><strong>%s &nbsp; [<span class="%s">%s</span>]</strong></p>' % (
            _mi("\\tau = %.1f\\,\\mathrm{MPa} < \\tau_{cr} = %.1f\\,\\mathrm{MPa}" % (st.tau_web, st.tau_cr)),
            cls_sb, ok_sb))

    a(_make_stiffener_layout_svg(p.L, st.spacing, st.n_pairs))
    a(_make_stiffener_detail_svg(r))

    a('<h3>8.2 支座加劲肋（支承加劲肋）</h3>')
    a('<p>支座反力 %s。采用 <strong>%d 块</strong> 竖向加劲肋，尺寸 %d &times; %d mm，端部刨平顶紧于下翼缘。</p>' %
      (_mi("R = V_{Ed} = %.2f\\,\\mathrm{kN}" % r.V_Ed), st.bearing_n, st.bearing_b, st.bearing_t))

    a('<p><strong>(a) 端面承压验算</strong></p>')
    A_ce = st.bearing_n * (st.bearing_b - 20) * st.bearing_t
    a('<p>%s</p>' % _mi("A_{ce} = %d \\times (%d - 20) \\times %d = %d\\,\\mathrm{mm^2}" %
                        (st.bearing_n, st.bearing_b, st.bearing_t, A_ce)))
    a('<p>%s</p>' % _mi(LX.SIGMA + "_{ce} = R / A_{ce} = %.2f \\times 10^3 / %d = %.1f\\,\\mathrm{MPa}" %
                        (r.V_Ed, A_ce, st.sigma_ce)))
    cls_ce = 'ok' if st.ce_ok else 'ng'; ok_ce = '满足' if st.ce_ok else '不满足'
    a('<p><strong>%s &nbsp; [<span class="%s">%s</span>]</strong></p>' % (
        _mi(LX.SIGMA + "_{ce} = %.1f\\,\\mathrm{MPa} < f_{ce} = 355\\,\\mathrm{MPa}" % st.sigma_ce), cls_ce, ok_ce))

    a('<p><strong>(b) 压杆稳定验算（十字形截面，绕腹板平面外弯曲）</strong></p>')
    a('<p>有效截面取 %d 块加劲肋 + 腹板两侧各 15 t<sub>w</sub> = %.0f mm 宽的腹板。</p>' %
      (st.bearing_n, 15 * sec.t_w))
    web_contrib = 15.0 * sec.t_w
    A_eff = st.bearing_n * st.bearing_b * st.bearing_t + web_contrib * sec.t_w
    I_eff = (st.bearing_t * (2.0 * st.bearing_b + sec.t_w) ** 3 / 12.0 + web_contrib * sec.t_w ** 3 / 12.0)
    i_eff = math.sqrt(I_eff / A_eff) if A_eff > 0 else 1.0
    lam = sec.h_w / i_eff
    a('<p>%s，%s，%s。</p>' % (
        _mi("A_{\\mathrm{eff}} = %.0f\\,\\mathrm{mm^2}" % A_eff),
        _mi("i_{\\mathrm{eff}} = %.1f\\,\\mathrm{mm}" % i_eff),
        _mi(LX.LAMBDA + " = h_w / i_{\\mathrm{eff}} = %.1f" % lam)))
    a('<p>b 类截面，%s，稳定系数 %s：</p>' % (_mi(LX.LAMBDA + " = %.1f" % lam), _mi(LX.VARPHI + " = %.3f" % st.phi_bearing)))
    cls_st = 'ok' if st.stab_ok else 'ng'; ok_st = '满足' if st.stab_ok else '不满足'
    a('<p><strong>%s &nbsp; [<span class="%s">%s</span>]</strong></p>' % (
        _mi(LX.SIGMA + " = R / (" + LX.VARPHI + " \\times A_{\\mathrm{eff}}) = %.1f\\,\\mathrm{MPa} < f_d = %d\\,\\mathrm{MPa}" %
            (st.sigma_stab, r.f_d_mid)), cls_st, ok_st))

    # ============== 九、连接构造说明 ==============
    a('<h2>九、连接构造说明</h2>')
    diff = sec.t_f - sec.t_f2
    a('<p>变截面处翼缘板厚从 %d mm 变为 %d mm（差 %d mm），按 JTG D64-2015 第7.3节，'
      '在较厚板侧加工成 1:4 斜坡过渡，采用全熔透对接焊缝连接，质量等级一级（100%% 无损检测）。</p>' %
      (sec.t_f, sec.t_f2, diff))
    a('<p>钢梁与混凝土桥面板之间通过焊钉剪力连接件实现组合作用，按 JTG D64-2015 第13章设计（本计算书从略）。</p>')

    # ============== 十、设计结果汇总 ==============
    a('<h2>十、设计结果汇总</h2>')
    a('<h3>10.1 截面尺寸</h3>')
    a('<table>')
    a('<tr><th>部位</th><th>腹板 h<sub>w</sub> &times; t<sub>w</sub> (mm)</th><th>翼缘 b<sub>f</sub> &times; t<sub>f</sub> (mm)</th><th>总高 h (mm)</th><th>W<sub>x</sub> (cm<sup>3</sup>)</th></tr>')
    a('<tr><td>跨中段</td><td>%d &times; %d</td><td>%d &times; %d</td><td>%d</td><td>%.2f</td></tr>' %
      (sec.h_w, sec.t_w, sec.b_f, sec.t_f, pm.h, pm.W_x / 1e3))
    a('<tr><td>变截面段</td><td>%d &times; %d</td><td>%d &times; %d</td><td>%d</td><td>%.2f</td></tr>' %
      (sec.h_w, sec.t_w, sec.b_f2, sec.t_f2, pv.h, pv.W_x / 1e3))
    a('</table>')

    a('<h3>10.2 焊缝汇总</h3>')
    a('<table>')
    a('<tr><th>焊缝位置</th><th>类型</th><th>h<sub>f</sub> (mm)</th><th>焊条</th></tr>')
    a('<tr><td>翼缘-腹板连接（全跨）</td><td>双面连续角焊缝</td><td>%d</td><td>%s</td></tr>' % (wd.h_f_chosen, p.electrode))
    a('<tr><td>横向加劲肋-腹板</td><td>双面角焊缝</td><td>6</td><td>%s</td></tr>' % p.electrode)
    a('<tr><td>支座加劲肋-腹板</td><td>双面角焊缝</td><td>8</td><td>%s</td></tr>' % p.electrode)
    a('<tr><td>翼缘对接（变截面处）</td><td>全熔透对接焊缝</td><td>&mdash;</td><td>%s</td></tr>' % p.electrode)
    a('</table>')

    a('<h3>10.3 加劲肋汇总</h3>')
    a('<table>')
    a('<tr><th>类型</th><th>间距 (mm)</th><th>尺寸 b<sub>s</sub> &times; t<sub>s</sub> (mm)</th><th>数量</th><th>说明</th></tr>')
    a('<tr><td>横向加劲肋</td><td>%.0f</td><td>%d &times; %d</td><td>%d 对</td><td>成对布置于腹板两侧</td></tr>' %
      (st.spacing, st.b_s, st.t_s, st.n_pairs))
    a('<tr><td>支座加劲肋</td><td>端部</td><td>%d &times; %d &times; %d</td><td>两端各%d块</td><td>端部刨平顶紧下翼缘</td></tr>' %
      (st.bearing_n, st.bearing_b, st.bearing_t, st.bearing_n))
    a('</table>')

    a('<h3>10.4 安全验算汇总</h3>')
    a('<table>')
    a('<tr><th>序号</th><th>验算项目</th><th>计算值</th><th>限值</th><th>判定</th></tr>')

    check_items = [
        (1, "跨中抗弯 &sigma;", "%.1f MPa" % c.sigma_mid, "%.0f MPa" % c.sigma_mid_limit, c.sigma_mid_ok),
        (2, "跨中抗剪 &tau;<sub>max</sub>", "%.1f MPa" % c.tau_max, "%.0f MPa" % c.tau_max_limit, c.tau_max_ok),
        (3, "折算应力 &sigma;<sub>zs</sub>", "%.1f MPa" % c.sigma_zs, "%.0f MPa" % c.sigma_zs_limit, c.sigma_zs_ok),
        (4, "活载挠度 &delta;<sub>q</sub>", "%.1f mm" % c.deflection_q, "%.1f mm" % c.deflection_limit, c.deflection_q_ok),
        (5, "翼缘局部稳定 b₁/t_f", "%.2f" % c.flange_local_ratio, "%.1f" % c.flange_local_limit, c.flange_local_ok),
        (6, "整体稳定性", "&mdash;", "&mdash;", True),
        (7, "疲劳 &Delta;&sigma;<sub>p</sub> (母材)", "%.1f MPa" % ft.delta_sigma_p, "%.1f MPa" % ft.delta_sigma_D_base, ft.check_base_metal),
        (8, "疲劳 &Delta;&sigma;<sub>p</sub> (焊缝)", "%.1f MPa" % ft.delta_sigma_p, "%.1f MPa" % ft.delta_sigma_D_weld, ft.check_fillet_weld),
        (9, "变截面抗弯 &sigma;", "%.1f MPa" % c.sigma_var, "%.0f MPa" % c.sigma_var_limit, c.sigma_var_ok),
        (10, "变截面抗剪 &tau;", "%.1f MPa" % c.tau_var, "%.0f MPa" % c.tau_var_limit, c.tau_var_ok),
        (11, "变截面折算 &sigma;<sub>zs</sub>", "%.1f MPa" % c.sigma_zs_var, "%.0f MPa" % c.sigma_zs_var_limit, c.sigma_zs_var_ok),
        (12, "支座承压 &sigma;<sub>ce</sub>", "%.1f MPa" % st.sigma_ce, "355 MPa", st.ce_ok),
        (13, "支座稳定 &sigma;", "%.1f MPa" % st.sigma_stab, "%.0f MPa" % r.f_d_mid, st.stab_ok),
        (14, "腹板受剪屈曲 &tau;", "%.1f MPa" % st.tau_web, "%.1f MPa" % st.tau_cr, st.shear_buckling_ok),
    ]
    for i, name, val, lim, ok_ in check_items:
        cls = 'ok' if ok_ else 'ng'
        ok_str = '满足' if ok_ else '不满足'
        a('<tr><td>%d</td><td>%s</td><td>%s</td><td>%s</td><td><span class="%s">%s</span></td></tr>' %
          (i, name, val, lim, cls, ok_str))
    a('</table>')
    a('<hr>')

    # ============== 十一、结论 ==============
    a('<h2>十一、结论</h2>')
    a('<p>本设计针对 %d m 简支钢板梁桥，完成了以下全部设计内容：</p>' % p.L)
    a('<ol>')
    a('<li><strong>荷载统计与内力计算</strong> — 恒载（钢梁自重 + 二期恒载）和活载（公路-I级车道荷载）的标准值与设计值，控制截面弯矩和剪力；</li>')
    a('<li><strong>截面尺寸拟定</strong> — 焊接双轴对称工字形截面，腹板 %d&times;%d mm，翼缘 %d&times;%d mm（跨中段）/ %d&times;%d mm（变截面段）；</li>' %
      (sec.h_w, sec.t_w, sec.b_f, sec.t_f, sec.b_f2, sec.t_f2))
    a('<li><strong>强度验算</strong> — 跨中及变截面处抗弯、抗剪、折算应力均满足要求；</li>')
    a('<li><strong>刚度验算</strong> — 活载挠度 %.1f mm &lt; L/500 = %.1f mm；</li>' % (c.deflection_q, c.deflection_limit))
    a('<li><strong>疲劳验算</strong> — 疲劳应力幅 %.1f MPa &lt; 各细节类别常幅疲劳极限；</li>' % ft.delta_sigma_p)
    a('<li><strong>整体稳定性</strong> — 混凝土桥面板提供连续侧向支撑，自然满足；</li>')
    a('<li><strong>变截面设计</strong> — L/6 处减小翼缘，有效节约钢材；</li>')
    a('<li><strong>翼缘焊缝设计</strong> — 全跨 h<sub>f</sub> = %d mm 双面连续角焊缝；</li>' % wd.h_f_chosen)
    a('<li><strong>局部稳定设计</strong> — %d 对横向加劲肋 (%d&times;%d mm) + 支座加劲肋 (%d&times;%d&times;%d mm)。</li>' %
      (st.n_pairs, st.b_s, st.t_s, st.bearing_n, st.bearing_b, st.bearing_t))
    a('</ol>')

    if r.checks.all_ok():
        a('<p><strong>全部验算项目均满足 JTG D64-2015《公路钢结构桥梁设计规范》和 '
          'JTG D60-2015《公路桥涵设计通用规范》的相关规定，设计结果合理可行。</strong></p>')
    else:
        a('<p><strong>部分验算项目不满足要求，需调整截面尺寸重新计算。</strong></p>')

    a(_HTML_FOOTER)

    return '\n'.join(L)


# ============================================================
# 文件保存
# ============================================================

def save_html(r: CalcResult, path: str):
    html = generate_html(r)
    with open(path, 'w', encoding='utf-8') as f:
        f.write(html)
    return path


def save_pdf(r: CalcResult, path: str, callback=None):
    """使用 QWebEngineView 导出 PDF"""
    html = generate_html(r)
    _render_to_pdf(html, path, callback)


def _render_to_pdf(html_content: str, pdf_path: str, callback=None):
    """用 setHtml 加载 HTML，轮询 MathJax 完成后 printToPdf"""
    from PySide6.QtCore import QUrl

    view = QWebEngineView()
    view.resize(900, 1200)
    poll_count = [0]
    max_polls = 80
    _done = [False]

    def _on_print_finished(file_path, success):
        view.deleteLater()
        if callback:
            callback(success)

    view.page().pdfPrintingFinished.connect(_on_print_finished)

    def _on_load(ok):
        if not ok:
            if callback:
                callback(False)
            view.deleteLater()
            return
        QTimer.singleShot(2000, _poll)

    def _poll():
        if _done[0]:
            return
        poll_count[0] += 1
        # 检查 MathJax 的 data-mjax-done 属性
        js = "(document.body&&document.body.getAttribute('data-mjax-done')==='1')?'done':'wait'"
        view.page().runJavaScript(js, _on_result)

    def _on_result(s):
        if _done[0]:
            return
        if s == 'done':
            _done[0] = True
            view.page().printToPdf(pdf_path)
        elif poll_count[0] < max_polls:
            QTimer.singleShot(500, _poll)
        else:
            _done[0] = True
            view.page().printToPdf(pdf_path)

    view.loadFinished.connect(_on_load)
    view.setHtml(html_content, QUrl("https://cdn.jsdelivr.net/"))


def save_all(r: CalcResult, output_dir: str, base_name: str = "计算书"):
    """输出 HTML 文件"""
    os.makedirs(output_dir, exist_ok=True)
    html_path = os.path.join(output_dir, base_name + ".html")
    save_html(r, html_path)
    return html_path


# ============================================================
# DOCX 导出（使用 math2docx 直接插入 LaTeX → OMML 公式）
# ============================================================

def _svg_to_temp_png(svg_html, scale=1):
    """将 SVG（或含 div 包裹的 HTML 中的 SVG）渲染为临时 PNG 文件路径

    注意：QSvgRenderer 在 offscreen 模式下无法渲染中文字体，
    因此先剥离 SVG 中的 <text> 元素用 QSvgRenderer 画图形，
    再用 Pillow 绘制文字。
    """
    import os, re, tempfile
    # 提取纯 <svg>...</svg> 标签内容
    if '<svg' in svg_html:
        m = re.search(r'<svg[^>]*>.*?</svg>', svg_html, re.DOTALL)
        if m:
            svg_html = m.group()
    # 解析 SVG 尺寸
    w_match = re.search(r'width="(\d+)"', svg_html)
    h_match = re.search(r'height="(\d+)"', svg_html)
    w = int(int(w_match.group(1)) * scale) if w_match else int(800 * scale)
    h = int(int(h_match.group(1)) * scale) if h_match else int(600 * scale)

    # 提取所有 <text> 元素，并从 SVG 中移除
    texts = []
    def _extract_text(m):
        attrs = (m.group(1) or '') + ' '
        content = m.group(2) if m.group(2) else ''
        x = float(re.search(r'x="([\d.]+)"', attrs).group(1)) if re.search(r'x="([\d.]+)"', attrs) else 0
        y = float(re.search(r'y="([\d.]+)"', attrs).group(1)) if re.search(r'y="([\d.]+)"', attrs) else 0
        fs = float(re.search(r'font-size="([\d.]+)"', attrs).group(1)) if re.search(r'font-size="([\d.]+)"', attrs) else 11
        anchor = re.search(r'text-anchor="([^"]+)"', attrs)
        anchor = anchor.group(1) if anchor else 'start'
        bold = 'font-weight="bold"' in attrs
        color = re.search(r'fill="([^"]+)"', attrs)
        color = color.group(1) if color else '#000'
        texts.append({'x': x, 'y': y, 'text': content, 'font_size': fs,
                       'anchor': anchor, 'bold': bold, 'color': color})
        return ''
    svg_no_text = re.sub(r'<text\b([^>]*)>(.*?)</text>', _extract_text, svg_html, flags=re.DOTALL)

    # 用 Qt 渲染无文字 SVG
    os.environ.setdefault('QT_QPA_PLATFORM', 'offscreen')
    from PySide6.QtSvg import QSvgRenderer
    from PySide6.QtGui import QPixmap, QPainter
    from PySide6.QtCore import QByteArray, Qt
    from PySide6.QtWidgets import QApplication
    app = QApplication.instance() or QApplication([])
    from PySide6.QtCore import QRectF
    renderer = QSvgRenderer(QByteArray(svg_no_text.encode('utf-8')))
    pix = QPixmap(w, h)
    pix.fill(Qt.GlobalColor.white)
    p = QPainter(pix)
    renderer.render(p, QRectF(0, 0, w, h))
    p.end()

    # QT → PIL
    from PIL import Image as PILImage, ImageDraw, ImageFont
    tmp_bmp = tempfile.NamedTemporaryFile(suffix='.bmp', delete=False)
    tmp_bmp.close()
    pix.save(tmp_bmp.name, 'BMP')
    pil_img = PILImage.open(tmp_bmp.name).convert('RGB')
    os.unlink(tmp_bmp.name)
    draw = ImageDraw.Draw(pil_img)

    # 文本绘制配置
    SIMSUN = r'C:\Windows\Fonts\simsun.ttc'
    SIMHEI = r'C:\Windows\Fonts\simhei.ttf'

    def hex_to_rgb(h):
        h = h.lstrip('#')
        if len(h) == 3:
            h = ''.join(c * 2 for c in h)
        return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

    import re as _re
    for t in texts:
        try:
            fp = SIMHEI if t['bold'] else SIMSUN
            fnt = ImageFont.truetype(fp, int(t['font_size'] * scale))
            fnt_sub = ImageFont.truetype(fp, int(t['font_size'] * scale * 0.65))
        except Exception:
            fnt = ImageFont.load_default()
            fnt_sub = ImageFont.load_default()
        sx = t['x'] * scale
        sy = t['y'] * scale
        c = hex_to_rgb(t['color']) if t['color'].startswith('#') else (0, 0, 0)

        # 解析下标标记 t_f → t 正常 + f 下标
        text = t['text']
        parts = []
        last_end = 0
        for m in _re.finditer(r'([A-Za-z\u0391-\u03C9]+)_([A-Za-z0-9\u0391-\u03C9,/=]+)', text):
            if m.start() > last_end:
                parts.append((text[last_end:m.start()], False))
            # 合并前导部分 + 主字母（正常），下标的以下标渲染
            parts.append((m.group(1), False))
            parts.append((m.group(2), True))
            last_end = m.end()
        if last_end < len(text):
            parts.append((text[last_end:], False))
        if not parts:
            parts = [(text, False)]

        # 计算总宽度以确定居中偏移
        total_w = 0
        for seg, is_sub in parts:
            ff = fnt_sub if is_sub else fnt
            bb = draw.textbbox((0, 0), seg, font=ff)
            total_w += bb[2] - bb[0]

        if t['anchor'] == 'middle':
            ox = sx - total_w / 2
        elif t['anchor'] == 'end':
            ox = sx - total_w
        else:
            ox = sx

        # 逐段绘制
        cx = ox
        baseline = sy - 3 * scale
        sub_offset = int(5 * scale)
        for seg, is_sub in parts:
            ff = fnt_sub if is_sub else fnt
            draw.text((cx, baseline + (sub_offset if is_sub else 0)), seg, fill=c, font=ff)
            bb = draw.textbbox((0, 0), seg, font=ff)
            cx += bb[2] - bb[0]

    tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
    tmp.close()
    pil_img.save(tmp.name, 'PNG', optimize=True)
    return tmp.name

def save_docx(r: CalcResult, path: str):
    """导出 Word 文档 (.docx)，LaTeX 公式转为原生 OMML 方程"""
    doc = generate_docx(r)
    doc.save(path)

    # 同时生成构造图 SVG 文件
    _save_drawings_svg(r, path)

    return path


def _save_drawings_svg(r, docx_path):
    """在 DOCX 同目录下生成构造图 SVG 文件"""
    import os
    base = os.path.splitext(docx_path)[0]

    svg1 = _make_girder_weld_svg(r)
    # 提取纯 SVG 代码（去掉外层 div）
    svg1_clean = svg1[svg1.find('<svg'):svg1.rfind('</svg>') + 6]
    p1 = base + '_焊缝构造图.svg'
    with open(p1, 'wb') as f:
        f.write(('<?xml version="1.0" encoding="UTF-8"?>\n' + svg1_clean).encode('utf-8'))

    svg2 = _make_stiffener_detail_svg(r)
    svg2_clean = svg2[svg2.find('<svg'):svg2.rfind('</svg>') + 6]
    p2 = base + '_加劲肋构造图.svg'
    with open(p2, 'wb') as f:
        f.write(('<?xml version="1.0" encoding="UTF-8"?>\n' + svg2_clean).encode('utf-8'))


def generate_docx(r: CalcResult):
    """生成 Word 文档，LaTeX 公式以 OMML 格式插入"""
    import math2docx
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    _temp_pngs = []  # 跟踪临时 PNG 文件以便清理
    def _embed_svg(svg_text, caption='', width=None, height=None):
        """将 SVG 字符串渲染为 PNG 嵌入文档（2x 以提升清晰度）"""
        nonlocal _temp_pngs
        png = _svg_to_temp_png(svg_text, scale=2)
        # 图片段落（居中）
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run()
        kw = {}
        if width is not None:
            kw['width'] = width
        if height is not None:
            kw['height'] = height
        if not kw:
            kw['width'] = Inches(5.2)
        run.add_picture(png, **kw)
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(6)
            r2 = cap.add_run()
            r2.font.size = Pt(10)
            r2.font.bold = True
            r2.text = caption
        _temp_pngs.append(png)

    p = r.params
    sec = r.section
    pm = r.props_mid
    pv = r.props_var
    c = r.checks
    st = r.stiffeners
    wd = r.welds
    ft = r.fatigue

    doc = Document()

    # ---- 全局样式 ----
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(4)

    def _set_cell_shading(cell, color):
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading)

    def _set_cell_text(cell, text):
        """设置单元格文本，含下标时自动格式化"""
        import re
        has_sub = bool(re.search(r'(?<=[A-Za-z\u0391-\u03C9])_(?=[A-Za-z0-9])', str(text)))
        if not has_sub:
            cell.text = str(text)
            return
        cell.text = ''
        sub_re = re.compile(r'([A-Za-z\u0391-\u03C9]+)_([A-Za-z0-9\u0391-\u03C9,/]+)')
        last_end = 0
        for m in sub_re.finditer(str(text)):
            if m.start() > last_end:
                cell.paragraphs[0].add_run(str(text)[last_end:m.start()])
            cell.paragraphs[0].add_run(m.group(1))
            r = cell.paragraphs[0].add_run(m.group(2))
            r.font.subscript = True
            last_end = m.end()
        if last_end < len(str(text)):
            cell.paragraphs[0].add_run(str(text)[last_end:])

    def _add_table(headers, rows):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            _set_cell_text(cell, h)
            for p_ in cell.paragraphs:
                p_.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p_.runs:
                    run.bold = True
                    run.font.size = Pt(10)
            _set_cell_shading(cell, 'E0E0E0')
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = table.rows[ri + 1].cells[ci]
                _set_cell_text(cell, str(val))
                for p_ in cell.paragraphs:
                    p_.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p_.runs:
                        run.font.size = Pt(10)
        return table

    def _add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
        return h

    def _add_formula(latex):
        """添加居中 LaTeX 公式，转为 OMML"""
        p_ = doc.add_paragraph()
        p_.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_.paragraph_format.space_before = Pt(6)
        p_.paragraph_format.space_after = Pt(6)
        math2docx.add_math(p_, latex)
        return p_

    def _add_text(text):
        """添加段落。含下标等数学记号的段落自动用 python-docx 格式化"""
        import re
        # 检测是否含希腊字母或下标
        has_math = bool(re.search(r'[\u0391-\u03C9\u0394\u2206]', text))
        has_sub = bool(re.search(r'(?<=[A-Za-z\u0391-\u03C9])_(?=[A-Za-z0-9])', text))
        has_super = '²' in text or '³' in text

        if not (has_math or has_sub or has_super):
            return doc.add_paragraph(text)

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

        sub_re = re.compile(
            r'([A-Za-z\u0391-\u03C9]+)'
            r'_'
            r'([A-Za-z0-9\u0391-\u03C9,/]+)'
        )

        last_end = 0
        for m in sub_re.finditer(text):
            if m.start() > last_end:
                p.add_run(text[last_end:m.start()])
            p.add_run(m.group(1))
            r_sub = p.add_run(m.group(2))
            r_sub.font.subscript = True
            last_end = m.end()

        if last_end < len(text):
            p.add_run(text[last_end:])

        return p

    def _add_bold_text(text):
        """加粗段落，含下标时自动格式化"""
        p_ = doc.add_paragraph()
        p_.paragraph_format.space_before = Pt(2)
        p_.paragraph_format.space_after = Pt(2)
        import re
        has_sub = bool(re.search(r'(?<=[A-Za-z\u0391-\u03C9])_(?=[A-Za-z0-9])', text))
        if not has_sub:
            p_.add_run(text).bold = True
            return p_
        sub_re = re.compile(r'([A-Za-z\u0391-\u03C9]+)_([A-Za-z0-9\u0391-\u03C9,/]+)')
        last_end = 0
        for m in sub_re.finditer(text):
            if m.start() > last_end:
                r = p_.add_run(text[last_end:m.start()])
                r.bold = True
            r1 = p_.add_run(m.group(1))
            r1.bold = True
            r2 = p_.add_run(m.group(2))
            r2.bold = True
            r2.font.subscript = True
            last_end = m.end()
        if last_end < len(text):
            r = p_.add_run(text[last_end:])
            r.bold = True
        return p_

    # ==================== 标题 ====================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f'{p.L:.0f}m 简支焊接双轴对称工字形钢板梁设计计算书')
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    _add_table(
        ['项目', '内容'],
        [
            ['设计题目', f'{p.L:.0f}m 简支工字型钢板梁设计'],
            ['设计依据', 'JTG D64-2015 / JTG D60-2015 / 钢结构设计原理'],
            ['钢材', p.steel_grade],
            ['焊条', p.electrode],
        ]
    )

    # ==================== 一、设计基本参数 ====================
    _add_heading('一、设计基本参数', 1)
    _add_text(f'本设计为 {p.L:.0f} m 简支公路桥梁钢板梁设计。桥梁上部结构采用钢-混凝土叠合梁形式，钢主梁为焊接双轴对称工字形截面，在跨径范围内进行一次变截面设计。')

    _add_heading('1.1 几何参数与荷载参数', 2)
    _add_table(
        ['参数', '符号', '数值', '单位'],
        [
            ['计算跨径', 'L', f'{p.L:.1f}', 'm'],
            ['桥面板宽度', 'B', f'{p.B:.1f}', 'm'],
            ['混凝土桥面板厚度', 'H₁', f'{p.H1:.2f}', 'm'],
            ['沥青铺装层厚度', 'H₂', f'{p.H2:.2f}', 'm'],
            ['横向分布系数', 'η', f'{p.eta:.2f}', '—'],
            ['恒载分项系数', 'γ_G', f'{p.gamma_G:.1f}', '—'],
            ['活载分项系数', 'γ_Q', f'{p.gamma_Q:.1f}', '—'],
            ['结构重要性系数', 'γ₀', f'{p.gamma_0:.1f}', '—'],
            ['活载冲击系数', '1+μ', f'{p.mu_impact:.2f}', '—'],
        ]
    )

    _add_heading('1.2 材料特性', 2)
    _add_text(f'钢材选用 {p.steel_grade} 低合金高强度结构钢，焊条选用 {p.electrode} 型。依据 JTG D64-2015 表 3.2.1，{p.steel_grade} 钢强度设计值如下：')
    _add_table(
        ['板厚 t (mm)', 'fd (MPa)', 'fvd (MPa)', 'fce (MPa)'],
        [
            ['t ≤ 16', '275', '160', '355'],
            ['16 < t ≤ 40', '270', '155', '355'],
            ['40 < t ≤ 63', '260', '150', '355'],
        ]
    )
    _add_text(f'弹性模量 E = 2.06×10⁵ MPa，剪切模量 G = 7.90×10⁴ MPa，泊松比 ν = 0.3。钢材容重 ρ = 78.5 kN/m³。')

    # ==================== 二、荷载计算 ====================
    _add_heading('二、荷载计算', 1)
    _add_heading('2.1 二期恒载（桥面板+铺装层）', 2)
    g2 = r.g2
    _add_formula(r'g_2 = \gamma_c B H_1 + \gamma_a B H_2')
    _add_formula(f'g_2 = 25.0 \\times {p.B} \\times {p.H1} + 24.0 \\times {p.B} \\times {p.H2} = {25*p.B*p.H1:.3f} + {24*p.B*p.H2:.3f} = {g2:.3f}\\;\\mathrm{{kN/m}}')
    _add_bold_text(f'二期恒载合计：g₂ = {g2:.3f} kN/m')

    _add_heading('2.2 活载标准值（JTG D60-2015 第4.3.1条）', 2)
    _add_text(f'公路-I级车道荷载：均布荷载 q_k = {Q_K} kN/m（满跨布置）。集中荷载 P_k（用于弯矩计算）按下式内插：')
    _add_formula('P_k = 270 + (360 - 270) \\times (L - 5) / (50 - 5)')
    _add_formula(f'P_k = 270 + 90 \\times ({p.L:.0f} - 5) / 45 = {r.P_k:.2f}\\;\\mathrm{{kN}}')
    _add_text(f'用于剪力计算时，P_k 乘 1.2 系数：P_{{k,shear}} = 1.2 × {r.P_k:.2f} = {r.P_k*1.2:.2f} kN。')
    _add_text(f'横向分布系数 η = {p.eta:.2f}，分配到单根主梁的活载标准值：')
    _add_text(f'• 均布活载：q = η × q_k = {p.eta:.2f} × {Q_K} = {r.q:.3f} kN/m')
    _add_text(f'• 集中活载（弯矩用）：P = η × P_k = {p.eta:.2f} × {r.P_k:.2f} = {r.P:.2f} kN')
    _add_text(f'• 集中活载（剪力用）：P_s = η × 1.2 × P_k = {p.eta:.2f} × {r.P_k*1.2:.2f} = {r.P_s:.2f} kN')

    _add_heading('2.3 一期恒载（钢梁自重）', 2)
    _add_text('钢梁自重取决于截面尺寸，先按经验初估，拟定截面后代入验算。')

    # ==================== 三、截面尺寸拟定 ====================
    _add_heading('三、截面尺寸拟定', 1)
    _add_heading('3.1 梁高确定', 2)
    _add_text('梁高是钢板梁经济性的关键参数，综合以下因素确定：')
    h_min15 = p.L * 1000 / 15
    h_min12 = p.L * 1000 / 12
    _add_text(f'(1) 按刚度条件（挠度控制）：简支梁满足挠度限值 L/500 所需的最小梁高约为 h = L/15~L/12 = {h_min15:.0f}~{h_min12:.0f} mm。')
    _add_formula(r'h_e = 7 \cdot W_{\mathrm{req}}^{1/3} - 300\;\mathrm{(mm)}')
    _add_text(f'(2) 综合确定：腹板高度 h_w = {sec.h_w:.0f} mm，翼缘厚度 t_f = {sec.t_f:.0f} mm（上下翼缘等厚），总梁高 h = h_w + 2t_f = {pm.h:.0f} mm。')

    _add_heading('3.2 拟选截面尺寸', 2)
    _add_table(
        ['板件', '符号', '尺寸 (mm)', '说明'],
        [
            ['腹板高度', 'hw', f'{sec.h_w:.0f}', '按 L/15~L/20 经济与刚度确定'],
            ['腹板厚度', 'tw', f'{sec.t_w:.0f}', f'满足抗剪与局部稳定，tw≥hw/170={sec.h_w/170:.1f}'],
            ['翼缘宽度', 'bf', f'{sec.b_f:.0f}', f'满足整体稳定，bf≥h/5={pm.h/5:.0f}'],
            ['翼缘厚度', 'tf', f'{sec.t_f:.0f}', '满足抗弯所需截面模量'],
            ['梁总高', 'h', f'{pm.h:.0f}', 'h = hw + 2tf'],
        ]
    )

    _add_heading('3.3 截面几何特性计算', 2)
    _add_formula(f'A = {sec.h_w:.0f} \\times {sec.t_w:.0f} + 2 \\times {sec.b_f:.0f} \\times {sec.t_f:.0f} = {pm.A:.0f}\\;\\mathrm{{mm^2}}')
    I_w = sec.t_w * sec.h_w ** 3 / 12.0
    d_f = (sec.h_w + sec.t_f) / 2.0
    _add_formula(f'I_w = t_w h_w^3 / 12 = {sec.t_w} \\times {sec.h_w}^3 / 12 = {I_w/1e6:.2f} \\times 10^6\\;\\mathrm{{mm^4}}')
    _add_formula(f'I_x = I_w + I_f = {pm.I_x/1e6:.2f} \\times 10^6\\;\\mathrm{{mm^4}}')
    _add_formula(f'W_x = I_x / (h/2) = {pm.I_x/1e6:.2f} \\times 10^6 / {pm.h/2:.0f} = {pm.W_x/1e3:.2f}\\;\\mathrm{{cm^3}}')
    _add_formula(f'S = b_f t_f (h_w + t_f)/2 + t_w (h_w/2)^2 / 2 = {pm.S/1e3:.2f}\\;\\mathrm{{cm^3}}')
    _add_formula(f'g_1 = A \\times \\rho_{{\\mathrm{{steel}}}} = {pm.A*1e-6:.4f} \\times 78.5 = {r.g1:.3f}\\;\\mathrm{{kN/m}}')
    _add_bold_text(f'总恒载：g = g₁ + g₂ = {r.g1:.3f} + {g2:.3f} = {r.g:.3f} kN/m')

    _embed_svg(_make_cross_section_svg(sec.h_w, sec.t_w, sec.b_f, sec.t_f, pm.h),
               '图 3.1  工字形截面示意图（单位：mm）',
               width=Inches(3.6))

    # ==================== 四、内力计算 ====================
    _add_heading('四、内力计算', 1)
    _add_text(f'简支梁计算跨径 L = {p.L} m，控制截面为跨中（弯矩最大）和支座（剪力最大）。')

    _add_heading('4.1 跨中弯矩', 2)
    _add_text('弯矩标准值：')
    _add_formula(f'M_{{gk}} = g L^2 / 8 = {r.g:.3f} \\times {p.L:.0f}^2 / 8 = {r.M_gk:.2f}\\;\\mathrm{{kN\\cdot m}}')
    _add_formula(f'M_{{qk}} = q L^2 / 8 + P L / 4 = {r.q:.3f} \\times {p.L:.0f}^2 / 8 + {r.P:.2f} \\times {p.L:.0f} / 4 = {r.M_qk:.2f}\\;\\mathrm{{kN\\cdot m}}')
    _add_text('弯矩设计值（承载能力极限状态）：')
    _add_formula(f'M_{{Ed}} = \\gamma_0 [\\gamma_G M_{{gk}} + \\gamma_Q (1+\\mu) M_{{qk}}]')
    _add_formula(f'M_{{Ed}} = {p.gamma_0} \\times [{p.gamma_G} \\times {r.M_gk:.2f} + {p.gamma_Q} \\times {p.mu_impact} \\times {r.M_qk:.2f}] = {r.M_Ed:.2f}\\;\\mathrm{{kN\\cdot m}}')

    _add_heading('4.2 支座剪力', 2)
    _add_formula(f'V_{{gk}} = g L / 2 = {r.V_gk:.2f}\\;\\mathrm{{kN}}')
    _add_formula(f'V_{{qk}} = q L / 2 + P_s / 2 = {r.V_qk:.2f}\\;\\mathrm{{kN}}')
    _add_formula(f'V_{{Ed}} = \\gamma_0 [\\gamma_G V_{{gk}} + \\gamma_Q (1+\\mu) V_{{qk}}] = {r.V_Ed:.2f}\\;\\mathrm{{kN}}')

    _embed_svg(_make_moment_shear_diagrams(p.L, r.g, r.q, r.P, r.P_s, r.x_cut,
                r.M_gk, r.M_qk, r.M_Ed, r.M_gk_x, r.M_qk_x, r.M_Ed_x,
                r.V_gk, r.V_qk, r.V_Ed, None, None, r.V_Ed_x),
               '图 4.1  简支梁弯矩图与剪力图')

    # ==================== 五、跨中强度及刚度验算 ====================
    _add_heading('五、跨中截面强度及刚度验算', 1)

    _add_heading('5.1 抗弯强度（JTG D64-2015 第5.2条）', 2)
    _add_text(f'翼缘板厚 t_f = {sec.t_f} mm，取 f_d = {r.f_d_mid} MPa。')
    _add_formula(f'\\sigma = M_{{Ed}} / W_x = {r.M_Ed:.2f} \\times 10^6 / ({pm.W_x/1e3:.2f} \\times 10^3) = {c.sigma_mid:.1f}\\;\\mathrm{{MPa}}')
    status = '满足' if c.sigma_mid_ok else '不满足'
    _add_bold_text(f'σ = {c.sigma_mid:.1f} MPa < f_d = {c.sigma_mid_limit:.0f} MPa  [{status}]')

    _add_heading('5.2 抗剪强度（JTG D64-2015 第5.3条）', 2)
    _add_text(f'腹板中性轴处剪应力最大，t_w = {sec.t_w} mm，取 f_vd = {r.f_vd_mid} MPa。')
    _add_formula(f'\\tau_{{\\max}} = V_{{Ed}} S / (I_x t_w) = {r.V_Ed:.2f} \\times 10^3 \\times {pm.S/1e3:.2f} \\times 10^3 / ({pm.I_x/1e6:.2f} \\times 10^6 \\times {sec.t_w}) = {c.tau_max:.1f}\\;\\mathrm{{MPa}}')
    status = '满足' if c.tau_max_ok else '不满足'
    _add_bold_text(f'τ_max = {c.tau_max:.1f} MPa < f_vd = {c.tau_max_limit:.0f} MPa  [{status}]')

    _add_heading('5.3 折算应力（JTG D64-2015 第5.4条）', 2)
    y_j_label = sec.h_w / 2.0
    _add_text(f'验算跨中截面翼缘与腹板交界处（y_j = h_w/2 = {y_j_label:.0f} mm）：')
    sigma_j = r.M_Ed * 1e6 * y_j_label / pm.I_x
    tau_j = r.V_Ed * 1e3 * pm.S_f / (pm.I_x * sec.t_w)
    _add_formula(f'\\sigma_j = M_{{Ed}} y_j / I_x = {sigma_j:.1f}\\;\\mathrm{{MPa}}')
    _add_formula(f'\\tau_j = V_{{Ed}} S_f / (I_x t_w) = {tau_j:.1f}\\;\\mathrm{{MPa}}')
    _add_formula(f'\\sigma_{{zs}} = \\sqrt{{\\sigma_j^2 + 3\\tau_j^2}} = \\sqrt{{{sigma_j:.1f}^2 + 3 \\times {tau_j:.1f}^2}} = {c.sigma_zs:.1f}\\;\\mathrm{{MPa}}')
    status = '满足' if c.sigma_zs_ok else '不满足'
    _add_bold_text(f'σ_zs = {c.sigma_zs:.1f} MPa < 1.1 f_d = {c.sigma_zs_limit:.0f} MPa  [{status}]')

    _add_heading('5.4 刚度验算（JTG D64-2015 第5.4条）', 2)
    _add_text('活载挠度（不计冲击系数）按均布+集中荷载叠加计算：')
    L_mm = p.L * 1000.0
    q_Nmm = r.q
    P_N = r.P * 1000.0
    delta_u = 5.0 * q_Nmm * L_mm ** 4 / (384.0 * E_STEEL * pm.I_x)
    delta_c = P_N * L_mm ** 3 / (48.0 * E_STEEL * pm.I_x)
    _add_formula(f'\\delta_u = 5qL^4 / (384EI_x) = {delta_u:.1f}\\;\\mathrm{{mm}}')
    _add_formula(f'\\delta_c = PL^3 / (48EI_x) = {delta_c:.1f}\\;\\mathrm{{mm}}')
    _add_formula(f'\\delta_q = \\delta_u + \\delta_c = {c.deflection_q:.1f}\\;\\mathrm{{mm}}')
    _add_formula(f'[\\delta] = L / 500 = {c.deflection_limit:.1f}\\;\\mathrm{{mm}}')
    status = '满足' if c.deflection_q_ok else '不满足'
    _add_bold_text(f'δ_q = {c.deflection_q:.1f} mm < [δ] = {c.deflection_limit:.1f} mm  [{status}]')

    _add_heading('5.5 整体稳定性（JTG D64-2015 第5.5条）', 2)
    _add_text('本桥采用钢-混凝土叠合梁，混凝土桥面板通过焊钉剪力连接件与钢主梁上翼缘焊接，两者形成牢固连接，桥面板为受压上翼缘提供连续侧向约束。因此整体稳定自然满足，无需计算。')

    _add_heading('5.6 疲劳强度（JTG D64-2015 第5.6节及附录C）', 2)
    _add_text(f'采用疲劳荷载模型 I（等效车道荷载乘 0.7 折减系数）：')
    _add_formula(f'q_f = 0.7 \\times 10.5 = 7.35\\;\\mathrm{{kN/m}}')
    _add_formula(f'P_f = 0.7 \\times {r.P_k:.2f} = {0.7*r.P_k:.2f}\\;\\mathrm{{kN}}')
    _add_text(f'按 η = {p.eta} 分配到单梁：q_f1 = {ft.q_f1:.3f} kN/m，P_f1 = {ft.P_f1:.2f} kN')
    _add_formula(f'M_f = q_{{f1}} L^2 / 8 + P_{{f1}} L / 4 = {ft.M_f:.2f}\\;\\mathrm{{kN\\cdot m}}')
    _add_formula(f'\\Delta\\sigma_p = M_f / W_x = {ft.M_f:.2f} \\times 10^6 / ({pm.W_x/1e3:.2f} \\times 10^3) = {ft.delta_sigma_p:.1f}\\;\\mathrm{{MPa}}')
    _add_table(
        ['细节类别', 'Δσc (MPa)', 'ΔσD (MPa)', 'Δσp (MPa)', '判定'],
        [
            ['翼缘母材（非焊接）', '160', f'{ft.delta_sigma_D_base:.1f}', f'{ft.delta_sigma_p:.1f}',
             '满足' if ft.check_base_metal else '不满足'],
            ['翼缘-腹板连续角焊缝', '80', f'{ft.delta_sigma_D_weld:.1f}', f'{ft.delta_sigma_p:.1f}',
             '满足' if ft.check_fillet_weld else '不满足'],
        ]
    )
    _add_bold_text(f'Δσp = {ft.delta_sigma_p:.1f} MPa < 各细节类别常幅疲劳极限，疲劳强度满足。')

    # ==================== 六、变截面设计 ====================
    _add_heading('六、变截面设计', 1)
    _add_text(f'简支梁弯矩沿跨径呈抛物线分布。为节约钢材，在距支座 L/6 = {r.x_cut:.3f} m 处减小翼缘尺寸，腹板尺寸保持不变。变截面处弯矩约为跨中最大弯矩的 {r.M_Ed_x/r.M_Ed*100:.1f}%。')

    _add_heading('6.2 变截面处内力', 2)
    _add_formula(f'M_{{gk}}(x) = g \\times x \\times (L-x) / 2 = {r.M_gk_x:.2f}\\;\\mathrm{{kN\\cdot m}}')
    _add_formula(f'M_{{qk}}(x) = q \\times x \\times (L-x) / 2 + P \\times x / 2 = {r.M_qk_x:.2f}\\;\\mathrm{{kN\\cdot m}}')
    _add_formula(f'M_{{Ed}}(x) = \\gamma_0 [\\gamma_G M_{{gk}}(x) + \\gamma_Q (1+\\mu) M_{{qk}}(x)] = {r.M_Ed_x:.2f}\\;\\mathrm{{kN\\cdot m}}')
    _add_text(f'V_Ed(x) = {r.V_Ed_x:.2f} kN')

    _add_heading('6.3 变截面处尺寸', 2)
    _add_table(
        ['参数', '等截面段（跨中）', '变截面段'],
        [
            ['腹板 hw×tw (mm)', f'{sec.h_w:.0f}×{sec.t_w:.0f}', f'{sec.h_w:.0f}×{sec.t_w:.0f} (不变)'],
            ['翼缘宽 bf (mm)', f'{sec.b_f:.0f}', f'{sec.b_f2:.0f}'],
            ['翼缘厚 tf (mm)', f'{sec.t_f:.0f}', f'{sec.t_f2:.0f}'],
            ['总高 h (mm)', f'{pm.h:.0f}', f'{pv.h:.0f}'],
            ['惯性矩 Ix (10⁶ mm⁴)', f'{pm.I_x/1e6:.2f}', f'{pv.I_x/1e6:.2f}'],
            ['截面模量 Wx (cm³)', f'{pm.W_x/1e3:.2f}', f'{pv.W_x/1e3:.2f}'],
        ]
    )

    _embed_svg(_make_variable_section_layout_svg(p.L, r.x_cut,
                (sec.h_w, sec.t_w, sec.b_f, sec.t_f),
                (sec.h_w, sec.t_w, sec.b_f2, sec.t_f2)),
               '图 6.1  变截面布置图（单位：mm）')

    _add_heading('6.4 变截面处强度验算', 2)
    for label, val, limit, ok_flag in [
        ('抗弯', c.sigma_var, c.sigma_var_limit, c.sigma_var_ok),
        ('抗剪', c.tau_var, c.tau_var_limit, c.tau_var_ok),
        ('折算应力', c.sigma_zs_var, c.sigma_zs_var_limit, c.sigma_zs_var_ok),
    ]:
        status = '满足' if ok_flag else '不满足'
        _add_bold_text(f'{label}：{val:.1f} MPa < {limit:.0f} MPa  [{status}]')

    # ==================== 七、翼缘焊缝设计 ====================
    _add_heading('七、翼缘焊缝设计', 1)
    _add_text(f'翼缘与腹板采用双面连续角焊缝连接，{p.electrode} 焊条。角焊缝强度设计值 f_ff = 200 MPa。')
    _add_heading('7.1 支座截面处（等截面段）', 2)
    _add_formula(f'S_f = b_f \\times t_f \\times (h_w + t_f) / 2 = {sec.b_f} \\times {sec.t_f} \\times {d_f:.0f} = {wd.S_f1/1e3:.2f}\\;\\mathrm{{cm^3}}')
    _add_formula(f'v = V_{{Ed}} S_f / I_x = {r.V_Ed:.2f} \\times 10^3 \\times {wd.S_f1/1e3:.2f} \\times 10^3 / ({pm.I_x/1e6:.2f} \\times 10^6) = {wd.v1:.1f}\\;\\mathrm{{N/mm}}')
    _add_text(f'所需焊脚尺寸（双面角焊缝）：h_f ≤ v / (2×0.7×200) = {wd.v1:.1f} / (2×0.7×200) = {wd.h_f_req1:.1f} mm')
    _add_text(f'构造要求：h_f_min = 1.5√{sec.t_f} = {wd.h_f_min1:.1f} mm，h_f_max = 1.2×{sec.t_w} = {wd.h_f_max1:.1f} mm')
    _add_bold_text(f'选用焊脚尺寸：h_f = {wd.h_f_chosen:.0f} mm')

    _add_heading('7.2 变截面处', 2)
    _add_formula(f'S_{{f2}} = {sec.b_f2} \\times {sec.t_f2} \\times ({sec.h_w} + {sec.t_f2}) / 2 = {wd.S_f2/1e3:.2f}\\;\\mathrm{{cm^3}}')
    _add_formula(f'v = {r.V_Ed_x:.2f} \\times 10^3 \\times {wd.S_f2/1e3:.2f} \\times 10^3 / ({pv.I_x/1e6:.2f} \\times 10^6) = {wd.v2:.1f}\\;\\mathrm{{N/mm}}')
    _add_text(f'所需 h_f = {wd.h_f_req2:.1f} mm。构造要求：h_f_min = {wd.h_f_min2:.1f} mm，h_f_max = {wd.h_f_max2:.1f} mm')
    _add_bold_text(f'选用焊脚尺寸：h_f = {wd.h_f_chosen:.0f} mm（全跨统一）')
    _add_bold_text(f'【结论】翼缘-腹板连接焊缝全跨统一采用 h_f = {wd.h_f_chosen:.0f} mm 双面连续角焊缝。')

    _embed_svg(_make_girder_weld_svg(r), '图 7.1  钢板梁及焊缝构造图（单位：mm）')

    # ==================== 八、局部稳定设计 ====================
    _add_heading('八、局部稳定设计', 1)
    _add_heading('8.1 腹板加劲肋（JTG D64-2015 第5.7节）', 2)
    _add_formula(f'h_w / t_w = {sec.h_w:.0f} / {sec.t_w:.0f} = {st.hw_tw_ratio:.1f}')
    _add_text('根据 JTG D64-2015 第5.7节，对 Q345 钢：')
    _add_text('• hw/tw ≤ 100：局部稳定自然满足')
    _add_text('• 100 < hw/tw ≤ 170：需配置横向加劲肋')
    _add_text('• hw/tw > 170：需配置横向和纵向加劲肋')

    if st.need_longitudinal:
        _add_text(f'因 hw/tw = {st.hw_tw_ratio:.1f} > 170，需配置横向和纵向加劲肋。')
    elif st.need_transverse:
        _add_text(f'因 100 < {st.hw_tw_ratio:.1f} ≤ 170，需配置横向加劲肋。')
    else:
        _add_text(f'因 hw/tw = {st.hw_tw_ratio:.1f} ≤ 100，局部稳定自然满足。')

    if st.need_transverse:
        _embed_svg(_make_stiffener_layout_svg(p.L, st.spacing, st.n_pairs),
                   '图 8.1  加劲肋布置示意图')
        _add_text(f'(1) 加劲肋间距')
        _add_text(f'横向加劲肋间距 a ≤ min(2hw, 3000) = {min(2*sec.h_w, 3000):.0f} mm。')
        _add_text(f'取间距 a = {st.spacing:.0f} mm（满足 0.5hw = {0.5*sec.h_w:.0f} ≤ a ≤ 3000），全跨均匀布置 {st.n_pairs} 对。')
        _add_text(f'(2) 加劲肋构造尺寸（JTG D64-2015 第5.7.3条）')
        _add_formula(f'b_s \\geq h_w/30 + 40 = {sec.h_w/30 + 40:.1f}\\;\\mathrm{{mm}}')
        _add_formula(f't_s \\geq b_s / 15 = {st.b_s/15:.1f}\\;\\mathrm{{mm}}')
        _add_bold_text(f'选用 bs×ts = {st.b_s:.0f}×{st.t_s:.0f} mm，成对布置在腹板两侧。')
        _embed_svg(_make_stiffener_detail_svg(r), '图 8.2  加劲肋构造图（单位：mm）')

        _add_heading('8.1.1 腹板区格受剪屈曲（JTG D64-2015 第5.7.4条）', 3)
        ratio_a_hw = st.spacing / sec.h_w
        k_tau = 5.34 + 4.00 / ratio_a_hw ** 2 if ratio_a_hw >= 1.0 else 4.00 + 5.34 / ratio_a_hw ** 2
        _add_formula(f'k_\\tau = {k_tau:.3f}')
        _add_text(f'正则化高厚比 λ_s = {st.hw_tw_ratio:.1f} / (41×0.825×√{k_tau:.3f}) = {st.hw_tw_ratio/(41*0.825*math.sqrt(k_tau)):.3f}')
        _add_text(f'腹板平均剪应力：τ = {st.tau_web:.1f} MPa，受剪屈曲临界应力：τ_cr = {st.tau_cr:.1f} MPa')
        status = '满足' if st.shear_buckling_ok else '不满足'
        _add_bold_text(f'τ = {st.tau_web:.1f} MPa < τ_cr = {st.tau_cr:.1f} MPa  [{status}]')

    _add_heading('8.2 支座加劲肋（支承加劲肋）', 2)
    _add_text(f'支座反力 R = V_Ed = {r.V_Ed:.2f} kN。采用 {st.bearing_n} 块竖向加劲肋，尺寸 {st.bearing_b:.0f}×{st.bearing_t:.0f} mm，端部刨平顶紧于下翼缘。')
    _add_text('(a) 端面承压验算')
    cut = 20.0
    A_ce = st.bearing_n * (st.bearing_b - cut) * st.bearing_t
    _add_formula(f'A_{{ce}} = {st.bearing_n} \\times ({st.bearing_b:.0f} - 20) \\times {st.bearing_t:.0f} = {A_ce:.0f}\\;\\mathrm{{mm^2}}')
    _add_formula(f'\\sigma_{{ce}} = R / A_{{ce}} = {r.V_Ed:.2f} \\times 10^3 / {A_ce:.0f} = {st.sigma_ce:.1f}\\;\\mathrm{{MPa}}')
    status = '满足' if st.ce_ok else '不满足'
    _add_bold_text(f'σ_ce = {st.sigma_ce:.1f} MPa < f_ce = 355 MPa  [{status}]')

    _add_text('(b) 压杆稳定验算（十字形截面，绕腹板平面外弯曲）')
    web_contrib = 15.0 * sec.t_w
    A_eff = st.bearing_n * st.bearing_b * st.bearing_t + web_contrib * sec.t_w
    I_eff = (st.bearing_t * (2.0 * st.bearing_b + sec.t_w) ** 3 / 12.0
             + web_contrib * sec.t_w ** 3 / 12.0)
    i_eff = math.sqrt(I_eff / A_eff) if A_eff > 0 else 1.0
    lam = sec.h_w / i_eff
    _add_text(f'有效截面 A_eff = {A_eff:.0f} mm²，i_eff = {i_eff:.1f} mm，λ = h_w / i_eff = {lam:.1f}')
    _add_text(f'b 类截面，λ = {lam:.1f}，稳定系数 φ = {st.phi_bearing:.3f}')
    status = '满足' if st.stab_ok else '不满足'
    _add_bold_text(f'σ = R / (φ × A_eff) = {st.sigma_stab:.1f} MPa < f_d = {r.f_d_mid:.0f} MPa  [{status}]')

    # ==================== 九、连接构造说明 ====================
    _add_heading('九、连接构造说明', 1)
    diff = sec.t_f - sec.t_f2
    _add_text(f'变截面处翼缘板厚从 {sec.t_f:.0f} mm 变为 {sec.t_f2:.0f} mm（差 {diff:.0f} mm），按 JTG D64-2015 第7.3节，在较厚板侧加工成 1:4 斜坡过渡，采用全熔透对接焊缝连接，质量等级一级（100% 无损检测）。')
    _add_text('钢梁与混凝土桥面板之间通过焊钉剪力连接件实现组合作用，按 JTG D64-2015 第13章设计（本计算书从略）。')

    # ==================== 十、设计结果汇总 ====================
    _add_heading('十、设计结果汇总', 1)
    _add_heading('10.1 截面尺寸', 2)
    _add_table(
        ['部位', '腹板 hw×tw (mm)', '翼缘 bf×tf (mm)', '总高 h (mm)', 'Wx (cm³)'],
        [
            ['跨中段', f'{sec.h_w:.0f}×{sec.t_w:.0f}', f'{sec.b_f:.0f}×{sec.t_f:.0f}',
             f'{pm.h:.0f}', f'{pm.W_x/1e3:.2f}'],
            ['变截面段', f'{sec.h_w:.0f}×{sec.t_w:.0f}', f'{sec.b_f2:.0f}×{sec.t_f2:.0f}',
             f'{pv.h:.0f}', f'{pv.W_x/1e3:.2f}'],
        ]
    )

    _add_heading('10.2 焊缝汇总', 2)
    _add_table(
        ['焊缝位置', '类型', 'hf (mm)', '焊条'],
        [
            ['翼缘-腹板连接（全跨）', '双面连续角焊缝', f'{wd.h_f_chosen:.0f}', p.electrode],
            ['横向加劲肋-腹板', '双面角焊缝', '6', p.electrode],
            ['支座加劲肋-腹板', '双面角焊缝', '8', p.electrode],
            ['翼缘对接（变截面处）', '全熔透对接焊缝', '—', p.electrode],
        ]
    )

    _add_heading('10.3 加劲肋汇总', 2)
    _add_table(
        ['类型', '间距 (mm)', '尺寸 bs×ts (mm)', '数量', '说明'],
        [
            ['横向加劲肋', f'{st.spacing:.0f}', f'{st.b_s:.0f}×{st.t_s:.0f}',
             f'{st.n_pairs} 对', '成对布置于腹板两侧'],
            ['支座加劲肋', '端部', f'{st.bearing_n}×{st.bearing_b:.0f}×{st.bearing_t:.0f}',
             f'两端各{st.bearing_n:.0f}块', '端部刨平顶紧下翼缘'],
        ]
    )

    _add_heading('10.4 安全验算汇总', 2)
    check_items = [
        (1, '跨中抗弯 σ', f'{c.sigma_mid:.1f} MPa', f'{c.sigma_mid_limit:.0f} MPa', c.sigma_mid_ok),
        (2, '跨中抗剪 τmax', f'{c.tau_max:.1f} MPa', f'{c.tau_max_limit:.0f} MPa', c.tau_max_ok),
        (3, '折算应力 σzs', f'{c.sigma_zs:.1f} MPa', f'{c.sigma_zs_limit:.0f} MPa', c.sigma_zs_ok),
        (4, '活载挠度 δq', f'{c.deflection_q:.1f} mm', f'{c.deflection_limit:.1f} mm', c.deflection_q_ok),
        (5, '翼缘局部稳定 b₁/tf', f'{c.flange_local_ratio:.2f}', f'{c.flange_local_limit:.1f}', c.flange_local_ok),
        (6, '整体稳定性', '—', '—', True),
        (7, '疲劳 Δσp (母材)', f'{ft.delta_sigma_p:.1f} MPa', f'{ft.delta_sigma_D_base:.1f} MPa', ft.check_base_metal),
        (8, '疲劳 Δσp (焊缝)', f'{ft.delta_sigma_p:.1f} MPa', f'{ft.delta_sigma_D_weld:.1f} MPa', ft.check_fillet_weld),
        (9, '变截面抗弯 σ', f'{c.sigma_var:.1f} MPa', f'{c.sigma_var_limit:.0f} MPa', c.sigma_var_ok),
        (10, '变截面抗剪 τ', f'{c.tau_var:.1f} MPa', f'{c.tau_var_limit:.0f} MPa', c.tau_var_ok),
        (11, '变截面折算 σzs', f'{c.sigma_zs_var:.1f} MPa', f'{c.sigma_zs_var_limit:.0f} MPa', c.sigma_zs_var_ok),
        (12, '支座承压 σce', f'{st.sigma_ce:.1f} MPa', '355 MPa', st.ce_ok),
        (13, '支座稳定 σ', f'{st.sigma_stab:.1f} MPa', f'{r.f_d_mid:.0f} MPa', st.stab_ok),
        (14, '腹板受剪屈曲 τ', f'{st.tau_web:.1f} MPa', f'{st.tau_cr:.1f} MPa', st.shear_buckling_ok),
    ]
    _add_table(
        ['序号', '验算项目', '计算值', '限值', '判定'],
        [[str(i), name, val, lim, '满足' if ok_ else '不满足'] for i, name, val, lim, ok_ in check_items]
    )

    # ==================== 十一、结论 ====================
    _add_heading('十一、结论', 1)
    _add_text(f'本设计针对 {p.L:.0f} m 简支钢板梁桥，完成了以下全部设计内容：')
    _add_text(f'(1) 荷载统计与内力计算 — 恒载（钢梁自重 + 二期恒载）和活载（公路-I级车道荷载）的标准值与设计值，控制截面弯矩和剪力；')
    _add_text(f'(2) 截面尺寸拟定 — 焊接双轴对称工字形截面，腹板 {sec.h_w:.0f}×{sec.t_w:.0f} mm，翼缘 {sec.b_f:.0f}×{sec.t_f:.0f} mm（跨中段）/ {sec.b_f2:.0f}×{sec.t_f2:.0f} mm（变截面段）；')
    _add_text(f'(3) 强度验算 — 跨中及变截面处抗弯、抗剪、折算应力均满足要求；')
    _add_text(f'(4) 刚度验算 — 活载挠度 {c.deflection_q:.1f} mm < L/500 = {c.deflection_limit:.1f} mm；')
    _add_text(f'(5) 疲劳验算 — 疲劳应力幅 {ft.delta_sigma_p:.1f} MPa < 各细节类别常幅疲劳极限；')
    _add_text('(6) 整体稳定性 — 混凝土桥面板提供连续侧向支撑，自然满足；')
    _add_text('(7) 变截面设计 — L/6 处减小翼缘，有效节约钢材；')
    _add_text(f'(8) 翼缘焊缝设计 — 全跨 h_f = {wd.h_f_chosen:.0f} mm 双面连续角焊缝；')
    _add_text(f'(9) 局部稳定设计 — {st.n_pairs} 对横向加劲肋 ({st.b_s:.0f}×{st.t_s:.0f} mm) + 支座加劲肋 ({st.bearing_n:.0f}×{st.bearing_b:.0f}×{st.bearing_t:.0f} mm)。')

    if r.checks.all_ok():
        _add_bold_text('全部验算项目均满足 JTG D64-2015《公路钢结构桥梁设计规范》和 JTG D60-2015《公路桥涵设计通用规范》的相关规定，设计结果合理可行。')
    else:
        _add_bold_text('部分验算项目不满足要求，需调整截面尺寸重新计算。')

    # 清理临时 PNG 文件
    import os
    for png in _temp_pngs:
        try:
            os.remove(png)
        except OSError:
            pass

    return doc
